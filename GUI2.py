from cryptography.fernet import Fernet
import datetime
import string
import time
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
from docx import Document
from xlutils.copy import copy
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Cm
import xlrd

def patch_crypto_be_discovery():

    """
    Monkey patches cryptography's backend detection.
    Objective: support pyinstaller freezing.
    """

    from cryptography.hazmat import backends

    try:
        from cryptography.hazmat.backends.commoncrypto.backend import backend as be_cc
    except ImportError:
        be_cc = None

    try:
        from cryptography.hazmat.backends.openssl.backend import backend as be_ossl
    except ImportError:
        be_ossl = None

    backends._available_backends_list = [
        be for be in (be_cc, be_ossl) if be is not None
    ]

patch_crypto_be_discovery()

root = Tk()

root.title("Świadectwomat v0.31 - PL")
root.wm_minsize(920, 645)
root.iconbitmap(r'bin\tubes.ico')

listawezy = []
listaparametrow = ['0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0']

listaklientow = []

def ladujklientow():
    try:
        wbX = xlrd.open_workbook(
            '//SRV-DOKUMENTY/qc - kontrola jakości/QC_ŚWIADECTWA TUBES/ŚWIADECTWOMAT/spis/klient.xls',
            formatting_info=True)
    except FileNotFoundError:
        wbX = xlrd.open_workbook('bin/klient.xls', formatting_info=True)
        messagebox.showwarning('Twój INTERNET mówi NIE!',
                               'Nie można nawiązać połączenia ze spisem klientów na serwerze sieciowym!\n\nZaładowano listę klientów z zapasowego pliku lokalnego.\n\nUWAGA - funkcja dodania nowego klienta do listy jest WYŁĄCZONA!')
    global wsX
    wsX = wbX.sheet_by_index(0)
    listaklientow.clear()
    for row in wsX.col(0):
        key = b'Uf-FryHfa-JXDToRDClFsSrv2h56kzEUCK5imvQngbo='
        f = Fernet(key)
        row_value = row.value
        kodowanie = row_value.encode()
        odkryte = f.decrypt(kodowanie)
        odszyfrowane = odkryte.decode('utf-8')
        listaklientow.append(odszyfrowane)

ladujklientow()

def dodajklienta():
    try:
        wbnX = xlrd.open_workbook(
            '//SRV-DOKUMENTY/qc - kontrola jakości/QC_ŚWIADECTWA TUBES/ŚWIADECTWOMAT/spis/klient.xls',
            formatting_info=True)
        adresklienta = '//SRV-DOKUMENTY/qc - kontrola jakości/QC_ŚWIADECTWA TUBES/ŚWIADECTWOMAT/spis/klient.xls'
    except FileNotFoundError:
        messagebox.showerror('Zapis lokalny niemożliwy!',
                               'Dodanie klienta do bazy nie jest możliwe ze względu na brak połączenia ze spisem na serwerze sieciowym!')
        pass
    wsnX = wbnX.sheet_by_index(0)
    wbn1X = copy(wbnX)
    wsn1X = wbn1X.get_sheet(0)
    liczbawierszy = 0
    for _ in wsnX.col(0):
        liczbawierszy += 1
    dk1 = entryKlient1.get()
    dk2 = entryKlient2.get()
    dk3 = entryKlient3.get()
    key = b'Uf-FryHfa-JXDToRDClFsSrv2h56kzEUCK5imvQngbo='
    f = Fernet(key)
    dk1token = f.encrypt(bytes(dk1, 'utf-8'))
    dk2token = f.encrypt(bytes(dk2, 'utf-8'))
    dk3token = f.encrypt(bytes(dk3, 'utf-8'))
    dk11token = dk1token.decode('utf-8')
    dk22token = dk2token.decode('utf-8')
    dk33token = dk3token.decode('utf-8')
    daneklienta = [dk11token, dk22token, dk33token]
    potwierdzeniedodaniaklienta = messagebox.askyesno('Dodanie nowego klienta', 'Czy na pewno chcesz dodać następującego klienta:\n\n'+dk1+'\n'+dk2+'\n'+dk3+'\n\nPo dodaniu dane zostaną zaszyfrowane i nie będą łatwe do zedytowania.')
    if potwierdzeniedodaniaklienta:
        for col_index in range(0,3):
            wsn1X.write(liczbawierszy, col_index, daneklienta[col_index])
        wbn1X.save(adresklienta)
        messagebox.showinfo("Dodano klienta do listy!", "Dodano klienta do listy!\n\nGratuluję sukcesu!")
        ladujklientow()
    else:
        pass

def znajdzklienta():
    def wybierzklienta(event):
        wynikszukania = rezultatszukania.get(ANCHOR)
        listawybrana = wynikszukania.rsplit('     ')
        entryKlient1.delete(0, END)
        entryKlient1.insert(END, listawybrana[0])
        entryKlient2.delete(0, END)
        entryKlient2.insert(END, listawybrana[1])
        entryKlient3.delete(0, END)
        entryKlient3.insert(END, listawybrana[2])
        oknoklient.destroy()
        ladujklientow()
    oknoklient = Toplevel(root)
    oknoklient.iconbitmap(r'bin\tubes.ico')
    oknoklient.title('Klienci zawierający szukane słowo:')
    rezultatszukania = Listbox(oknoklient, height=10, activestyle='none', width=100)
    rezultatszukania.pack(side=LEFT, fill=BOTH)
    rezultatszukania.delete(0, END)
    rezultatszukania.bind("<Double-Button-1>", wybierzklienta)
    yscrollbarX = Scrollbar(oknoklient, orient=VERTICAL)
    yscrollbarX.config(command=rezultatszukania.yview)
    yscrollbarX.pack(side=RIGHT, fill=Y)
    rezultatszukania.config(yscrollcommand=yscrollbarX.set)
    wprowadzony = entryKlient1.get()
    listaponumerowana = list(enumerate(listaklientow))
    znalezioneindeksy = [i for i,x in listaponumerowana if wprowadzony.lower() in x.lower()]
    for item in znalezioneindeksy:
        wartoscadresu1 = wsX.cell(item,1).value
        key = b'Uf-FryHfa-JXDToRDClFsSrv2h56kzEUCK5imvQngbo='
        f = Fernet(key)
        kodowanie1 = wartoscadresu1.encode()
        odkryte1 = f.decrypt(kodowanie1)
        wartoscadresu11 = odkryte1.decode('utf-8')
        wartoscadresu2 = wsX.cell(item,2).value
        kodowanie2 = wartoscadresu2.encode()
        odkryte2 = f.decrypt(kodowanie2)
        wartoscadresu22 = odkryte2.decode('utf-8')
        rezultatszukania.insert(END, listaklientow[item]+'     '+wartoscadresu11+'     '+wartoscadresu22)

def znajdzklientaenter(event):
    znajdzklienta()

def ladujlistewezy():
    try:
        wb = xlrd.open_workbook(
            '//SRV-DOKUMENTY/qc - kontrola jakości/QC_ŚWIADECTWA TUBES/ŚWIADECTWOMAT/spis/spis.xls')
    except FileNotFoundError:
        wb = xlrd.open_workbook('bin/spis.xls')
        messagebox.showwarning('Nie można połączyć ze spisem węży',
                               'Nie udało się pobrać listy węży z folderu sieciowego. Załadowano listę z dysku lokalnego, która może zawierać nieaktualne dane.\nW razie kolejnych niepowodzeń skontaktuj się z Mariuszem!')
    global ws
    ws = wb.sheet_by_name('lista')
    listawezy.clear()
    for row in ws.col(0):
        listawezy.append(row.value)

ladujlistewezy()

def otworznowywaz():
    matka = Toplevel()

    matka.title('Dodaj opis przewodu')
    matka.iconbitmap(r'bin\tubes.ico')

    def dodajwazdolisty():
        try:
            wbn = xlrd.open_workbook('//SRV-DOKUMENTY/qc - kontrola jakości/QC_ŚWIADECTWA TUBES/ŚWIADECTWOMAT/spis/spis.xls', formatting_info=True)
            adresspisu = '//SRV-DOKUMENTY/qc - kontrola jakości/QC_ŚWIADECTWA TUBES/ŚWIADECTWOMAT/spis/spis.xls'
        except FileNotFoundError:
            wbn = xlrd.open_workbook('bin/spis.xls', formatting_info=True)
            messagebox.showwarning('Zapis lokalny!','Wąż dodany do listy lokalnej, zapasowej. Wprowadzone informacje o wężu NIE będą utrwalone w spisie sieciowym!')
            adresspisu = 'bin/spis.xls'
        wsn = wbn.sheet_by_name('lista')
        wbn1 = copy(wbn)
        wsn1 = wbn1.get_sheet(0)
        liczbawierszy = 0
        for _ in wsn.col(0):
            liczbawierszy += 1
        k1 = kolumnaindeks.get()
        k2 = kolumnatypweza.get()
        k3 = kolumnaopisweza.get()
        k4 = kolumnawarstwazewnetrzna.get()
        k5 = kolumnawzmocnienie.get()
        k6 = kolumnawarstwawewnetrzna.get()
        k7 = kolumnatempmaksymalna.get()
        k8 = kolumnatempminimalna.get()
        k9 = kolumnasrednica.get()
        k10 = kolumnacisnienierobocze.get()
        k11 = kolumnacisnienierozerwania.get()
        k12 = kolumnapromienstatyczny.get()
        k13 = kolumnapromiendynamiczny.get()
        k14 = kolumnaflagatolerancji.get()
        parametrydododania = [k1, k2, k3, k4, k5, k6, k7, k8, k9, k10, k11, k12, k13, k14]
        for col_index in range(0,14):
           wsn1.write(liczbawierszy, col_index, parametrydododania[col_index])
        wbn1.save(adresspisu)
        messagebox.showinfo("Dodano wąż do listy!", "Dodano wąż do listy!\n\nGratuluję sukcesu!")
        ladujlistewezy()
        matka.destroy()


    Label(matka, text='Indeks węża:').grid(padx=2, pady=2, sticky=W, row=0, column=0)

    kolumnaindeks = Entry(matka, width=30)
    kolumnaindeks.grid(padx=2, pady=2, sticky=E, row=0, column=1)

    Label(matka, text='Typ węża (np. 2SN, PARRAP, MP 20):').grid(padx=2, pady=2, sticky=W, row=1, column=0)

    kolumnatypweza = Entry(matka, width=30)
    kolumnatypweza.grid(padx=2, pady=2, sticky=E, row=1, column=1)

    Label(matka, text='Opis węża:').grid(padx=2, pady=2, sticky=W, row=2, column=0)

    kolumnaopisweza = Entry(matka, width=30)
    kolumnaopisweza.grid(padx=2, pady=2, sticky=E, row=2, column=1)

    Label(matka, text='Warstwa zewnętrzna:').grid(padx=2, pady=2, sticky=W, row=3, column=0)

    kolumnawarstwazewnetrzna = Entry(matka, width=30)
    kolumnawarstwazewnetrzna.grid(padx=2, pady=2, sticky=E, row=3, column=1)

    Label(matka, text='Wzmocnienie:').grid(padx=2, pady=2, sticky=W, row=4, column=0)

    kolumnawzmocnienie = Entry(matka, width=30)
    kolumnawzmocnienie.grid(padx=2, pady=2, sticky=E, row=4, column=1)

    Label(matka, text='Warstwa wewnętrzna:').grid(padx=2, pady=2, sticky=W, row=5, column=0)

    kolumnawarstwawewnetrzna = Entry(matka, width=30)
    kolumnawarstwawewnetrzna.grid(padx=2, pady=2, sticky=E, row=5, column=1)

    Label(matka, text='Temp. maksymalna (°C):').grid(padx=2, pady=2, sticky=W, row=6, column=0)

    kolumnatempmaksymalna = Entry(matka, width=30)
    kolumnatempmaksymalna.grid(padx=2, pady=2, sticky=E, row=6, column=1)

    Label(matka, text='Temp. minimalna (°C):').grid(padx=2, pady=2, sticky=W, row=7, column=0)

    kolumnatempminimalna = Entry(matka, width=30)
    kolumnatempminimalna.grid(padx=2, pady=2, sticky=E, row=7, column=1)

    Label(matka, text='Średnica wewnętrzna (mm):').grid(padx=2, pady=2, sticky=W, row=8, column=0)

    kolumnasrednica = Entry(matka, width=30)
    kolumnasrednica.grid(padx=2, pady=2, sticky=E, row=8, column=1)

    Label(matka, text='Ciśnienie robocze (bar):').grid(padx=2, pady=2, sticky=W, row=9, column=0)

    kolumnacisnienierobocze = Entry(matka, width=30)
    kolumnacisnienierobocze.grid(padx=2, pady=2, sticky=E, row=9, column=1)

    Label(matka, text='Ciśnienie rozerwania (bar):').grid(padx=2, pady=2, sticky=W, row=10, column=0)

    kolumnacisnienierozerwania = Entry(matka, width=30)
    kolumnacisnienierozerwania.grid(padx=2, pady=2, sticky=E, row=10, column=1)

    Label(matka, text='Promień zagięcia (statyczny) (mm):').grid(padx=2, pady=2, sticky=W, row=11, column=0)

    kolumnapromienstatyczny = Entry(matka, width=30)
    kolumnapromienstatyczny.grid(padx=2, pady=2, sticky=E, row=11, column=1)

    Label(matka, text='Promień zagięcia (dynamiczny) (mm):').grid(padx=2, pady=2, sticky=W, row=12, column=0)

    kolumnapromiendynamiczny = Entry(matka, width=30)
    kolumnapromiendynamiczny.grid(padx=2, pady=2, sticky=E, row=12, column=1)

    Label(matka, text='Flaga tolerancji:').grid(padx=2, pady=2, sticky=W, row=13, column=0)
    Button(matka, text='?', command=pomocflaga).grid(sticky=E, row=13, column=0)

    kolumnaflagatolerancji = ttk.Combobox(matka, state='readonly', width=27)
    kolumnaflagatolerancji['values'] = ['P1', 'P2', 'P3', 'S1', 'S2', 'S3', 'S4', 'S5', 'K', 'C']
    kolumnaflagatolerancji.grid(padx=2, pady=2, sticky=E, row=13, column=1)
    kolumnaflagatolerancji.current(0)

    Button(matka, text='Dodaj wąż!', height=2, command=dodajwazdolisty).grid(padx=2, pady=20, sticky=E, row=14, column=0, columnspan=1)

    matka.mainloop()

def pomocflaga():
    messagebox.showinfo('Pomoc nt. flag tolerancji', 'Flagi tolerancji to oznaczenie, dzięki któremu na świadectwie pojawia się informacja o tolerancji wykonania przewodu.\n\nKategorie:\nP - dla przewodów \"podstawowych\"\nS - dla przewodów stalowych\nK - dla przewodów kompozytowych\nC - dla przewodów Corroflon\n\nCyfry przy znacznikach P i S oznaczają kategorię ze względu na średnicę, tak jak podają instrukcje tolerancji dla tych rodzajów przewodów.\n\nW razie wątpliwości kontakt z Mariuszem.')

def sprawdzczywybranowaz():
    if indeksweza.get() != '':
        pass
    else:
        messagebox.showinfo("Nie wybrano indeksu węża!", "Nie wybrano indeksu węża!\n\nJeżeli szukanego węża nie ma liście, wybierz pustą pozycję z listy (tzw. zapchajdziurę)!")

def onselectoznakowanie(event):
    cowybrano = typoznaczenia.get()
    if cowybrano == 'NIE - oznaczenie standardowe':
        poleoznaczenia.delete(0.0, END)
    elif cowybrano == 'Pharmaline N/G':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'PHGP N DNXX PN'+cisnienieRob.get()+'\nTUBES INT. '+numerKRP1.get()+'/'+numerKRP2.get()+'/'+numerKRP3.get())
    elif cowybrano == 'CE I (stalowy)':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'TI-'+str(now.year)+'-EN14585-1-DNXX-PS'+cisnienieRob.get()+'\nTS TEMP°C / TEMP°C-'+numerKRP1.get()+'/'+numerKRP2.get()+'/'+numerKRP3.get())
    elif cowybrano == 'CE II (stalowy)':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'TI-'+str(now.year)+'-EN14585-1-DNXX-PS'+cisnienieRob.get()+'\nTS TEMP°C / TEMP°C-'+numerKRP1.get()+'/'+numerKRP2.get()+'/'+numerKRP3.get())
    elif cowybrano == 'CE I (niestalowy)':
        wyplujdane()
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'TUBES INTERNATIONAL\nDNXX '+listaparametrow[0]+ time.strftime(' %m/%y') + '\nWP/TP (BAR) ' + cisnienieRob.get() +'/'+cisnienieTest.get()+(' '+numerKRP1.get()+'/'+numerKRP2.get()+'/'+numerKRP3.get()))
    elif cowybrano == 'Spir Star':
        sprawdzczywybranowaz()
        wyplujdane()
        poleoznaczenia.delete(0.0, END)
        allow = string.digits
        poleoznaczenia.insert(1.0, 'TUBES '+time.strftime('%m %y'+'\n')+re.sub('[^%s]' % allow, '', numerKRP1.get()+'/'+numerKRP2.get()+'/'+numerKRP3.get())+'BNRXXXXX\nTYPE XX/XX WP '+str(int(listaparametrow[9]))+' BAR')

def szukaj(event):
    rezultaty.delete(0, END)
    wprowadzony = szukajweza.get()
    matching = [s for s in listawezy if wprowadzony.lower() in s.lower()]
    for item in matching:
        rezultaty.insert(END, item)

def wprowadzwybor():
    indeksweza.config(state='normal')
    indeksweza.delete(0, END)
    indeksweza.insert(END, rezultaty.get(ANCHOR))
    indeksweza.config(state='disabled')

def kliknijwybor(event):
    indeksweza.config(state='normal')
    indeksweza.delete(0, END)
    indeksweza.insert(END, rezultaty.get(ANCHOR))
    indeksweza.config(state='disabled')

def wyplujdane():
    wybranywaz = indeksweza.get()
    del listaparametrow[:]
    for wiersz in range(ws.nrows):
        if wybranywaz == ws.cell_value(wiersz,0):
            wierszweza = wiersz
            for cell in ws.row(wierszweza):
                cell_value = cell.value
                if cell.ctype in (2,3) and int(cell_value) == cell_value:
                    cell_value = str(int(cell_value))
                elif cell.ctype in (2,3) and str(cell_value) == cell_value:
                    cell_value = str(cell_value)
                listaparametrow.append(cell_value)

def wyczyscpola():
    potwierdzeniewyczyszczenia = messagebox.askyesno('Potwierdź', 'Czy chcesz wyczyścić wszystkie pola formularza?')
    if potwierdzeniewyczyszczenia:
        numerQC.delete(0, END)
        numerKRP1.delete(0, END)
        numerKRP2.delete(0, END)
        numerKRP3.delete(0, END)
        tu.delete(0, END)
        ilosc.delete(0, END)
        ilosc.insert(END, 1)
        dlugosc.delete(0, END)
        medium.delete(0, END)
        cisnienieRob.delete(0, END)
        tempWew.delete(0, END)
        tempZew.delete(0, END)
        mediumBadania.current([0])
        cisnienieTest.delete(0, END)
        czasTestu.delete(0, END)
        numerZam.delete(0, END)
        entryKlient1.delete(0, END)
        entryKlient2.delete(0, END)
        entryKlient3.delete(0, END)
        koncowka1entry.delete(0, END)
        koncowka2entry.delete(0, END)
        indeksweza.config(state='normal')
        indeksweza.delete(0, END)
        indeksweza.config(state='disabled')
        czyzuzycie.set(0)
        czyodtlu.set(0)
        czyciaglosc.set(0)
        szukajweza.delete(0, END)
        poleuwagi.delete('0.0', END)
        poleoznaczenia.delete('0.0', END)
        typoznaczenia.current(0)
    else:
        pass

def generatortolerancji():
    flaga = listaparametrow[13]
    dlu = int(dlugosc.get())
    global toler
    if flaga == 'P1':
        if dlu <= 630:
            toler = '-3 mm / +7 mm'
        elif 630 < dlu <= 1250:
            toler = '-4 mm / +12 mm'
        elif 1250 < dlu <= 2500:
            toler = '-6 mm / +20 mm'
        elif 2500 < dlu <= 8000:
            toler = '-0,5% / +1,5%'
        elif 8000 < dlu:
            toler = '-1% / +3%'
    elif flaga == 'P2':
        if dlu <= 630:
            toler = '-4 mm / +12 mm'
        elif 630 < dlu <= 1250:
            toler = '-6 mm / +20 mm'
        elif 1250 < dlu <= 2500:
            toler = '-6 mm / +25 mm'
        elif 2500 < dlu <= 8000:
            toler = '-0,5% / +1,5%'
        elif 8000 < dlu:
            toler = '-1% / +3%'
    elif flaga == 'P3':
        if dlu <= 2500:
            toler = '-6 mm / +25 mm'
        elif 2500 < dlu <= 8000:
            toler = '-0,5% / +1,5%'
        elif 8000 < dlu:
            toler = '-1% / +3%'
    elif flaga == 'S1':
        if dlu <= 1000:
            toler = '0 mm / +10 mm'
        elif dlu > 1000:
            toler = '0% / +1%'
    elif flaga == 'S2':
        if dlu <= 1000:
            toler = '0 mm / +15 mm'
        elif dlu > 1000:
            toler = '0% / +1,5%'
    elif flaga == 'S3':
        if dlu <= 1000:
            toler = '0 mm/ +20 mm'
        elif dlu > 1000:
            toler = '0% / +2%'
    elif flaga == 'S4':
        if dlu <= 1000:
            toler = '0 mm / +30 mm'
        elif dlu > 1000:
            toler = '0% / +3%'
    elif flaga == 'S5':
        if dlu <= 1000:
            toler = '0 mm / +40 mm'
        elif dlu > 1000:
            toler = '0% / +4%'
    elif flaga == 'K':
        if dlu <= 2500:
            toler = '± 50 mm'
        elif dlu > 2500:
            toler = '± 2%'
    elif flaga == 'C':
        if dlu <= 1000:
            toler = '0% / +5%'
        elif dlu > 1000:
            toler = '0% / +10%'
    return toler

def generujpopolsku():
    sprawdzczywybranowaz()
    wyplujdane()
    generatortolerancji()
    global checkIle, checkIle2
    nQC = numerQC.get()
    nKRP = numerKRP1.get()+'/'+numerKRP2.get()+'/'+numerKRP3.get()
    nTU = tu.get()
    ile = int(ilosc.get())
    rodzdlu = waz.get()
    dlu = dlugosc.get()
    med = medium.get()
    cisR = cisnienieRob.get()
    tw = tempWew.get()
    tz = tempZew.get()
    medB = mediumBadania.get()
    cisT = cisnienieTest.get()
    cz = czasTestu.get()
    zam = numerZam.get()
    przygotowal = entryprzygotowal.get()
    kiedy_p = entrykiedyprzygotowal.get()
    zatwierdzil = entryzatwierdził.get()
    kiedy_z = entrykiedyzatwierdzil.get()
    kli1 = entryKlient1.get()
    kli2 = entryKlient2.get()
    kli3 = entryKlient3.get()
    konc1 = koncowka1entry.get()
    konc2 = koncowka2entry.get()
    uwagi = poleuwagi.get(0.0, "end-1c")
    kol1 = listaparametrow[0]
    kol2 = listaparametrow[1]
    kol3 = listaparametrow[2]
    kol4 = listaparametrow[3]
    kol5 = listaparametrow[4]
    kol6 = listaparametrow[5]
    kol7 = listaparametrow[6]
    kol8 = listaparametrow[7]
    kol9 = listaparametrow[8]
    kol10 = listaparametrow[9]
    kol11 = listaparametrow[10]
    kol12 = listaparametrow[11]
    kol13 = listaparametrow[12]
    wyboroznak = typoznaczenia.get()

    if ile == 1:
        checkIle = 'przewód'
    elif ile > 4:
        checkIle = 'przewodów'
    elif ile == 2 or 3 or 4:
        checkIle = 'przewody'
    else:
        print('errorrrr')

    if ile == 1:
        checkIle2 = 'Przewód'
    elif ile > 1:
        checkIle2 = 'Przewody'
    else:
        print('errrororor2')

    if typoznaczenia.get() == 'CE I (stalowy)' or typoznaczenia.get() == 'CE II (stalowy)' or typoznaczenia.get() == 'CE I (niestalowy)':
        document = Document('bin\wzor_ce.docx')
        firstparagraph = document.paragraphs[0]
        firstparagraph.add_run('ŚWIADECTWO JAKOŚCI\tQC/'+nQC+time.strftime('/%y'))
        firstparagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firstparagraph.style = 'Heading 1'
    else:
        document = Document('bin\wzor2.docx')
        firstparagraph = document.paragraphs[0]
        firstparagraph.add_run('ŚWIADECTWO JAKOŚCI\tQC/' + nQC+time.strftime('/%y'))
        firstparagraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firstparagraph.style = 'Heading 1'

    klient = document.add_paragraph('Wystawiono dla:\t'+kli1)
    klient.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph('\t'+kli2).paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph('\t'+kli3).paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph()
    if zam == '':
        pass
    else:
        zamowienie = document.add_paragraph('Zamówienie:\t' + zam)
        zamowienie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()

    wykonanie = document.add_paragraph('Wykonanie:\t'+checkIle2+' wykonano wg Karty Roboczej Przewodu TI' + nKRP)
    wykonanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph()

    if wyboroznak == 'NIE - oznaczenie standardowe':
        oznakowanie = document.add_paragraph('Oznakowanie:\t'+checkIle2+' oznakowano TI' + nKRP)
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()
    elif wyboroznak == 'Pharmaline N/G':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t'+poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER)
        oznakowanie2 = document.add_paragraph('\t'+poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER)
        document.add_paragraph()
    elif wyboroznak == 'CE I (stalowy)':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        document.add_paragraph()
    elif wyboroznak == 'CE II (stalowy)':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        oznakowanie2 = document.add_paragraph('\t1433\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(10.25), WD_TAB_ALIGNMENT.LEFT)
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
    elif wyboroznak == 'CE I (niestalowy)':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        document.add_paragraph()
    elif wyboroznak == 'Spir Star':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER)
        document.add_paragraph()

    opisprzewodu = document.add_paragraph(
        'Wyrób:\t'+checkIle2+' TU/' + nTU + ' wykonano z węża typu '+kol2+' Ø '+str(kol9)+' mm ('+kol1+'), zakończono')
    opisprzewodu.paragraph_format.left_indent = (Cm(4))
    opisprzewodu.paragraph_format.first_line_indent = (Cm(-4))
    opisprzewodu.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    opisprzewodu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if ilekoncowek.get() == 1:
        opisprzewodu.add_run(' końcówkami:')
        koncowka1 = document.add_paragraph('Końcówka 1: '+konc1)
        koncowka1.paragraph_format.left_indent = (Cm(6.25))
        koncowka1.paragraph_format.first_line_indent = (Cm(-2.25))
        koncowka1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        koncowka2 = document.add_paragraph('Końcówka 2: '+konc2)
        koncowka2.paragraph_format.left_indent = (Cm(6.25))
        koncowka2.paragraph_format.first_line_indent = (Cm(-2.25))
        koncowka2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif ilekoncowek.get() == 2:
        opisprzewodu.add_run(' obustronnie końcówkami '+konc1+'.')
    else:
        pass

    zczegok2 = zczegok.get()
    zczegot2 = zczegot.get()
    if zczegok2 == 2 and zczegot2 == 2:
        materialkoncowek = document.add_paragraph('Końcówki i tuleje zostały wykonane ze stali nierdzewnej.')
        materialkoncowek.paragraph_format.left_indent = (Cm(4))
        materialkoncowek.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif zczegok2 == 1 and zczegot2 == 1:
        materialkoncowek = document.add_paragraph('Końcówki i tuleje zostały wykonane ze stali węglowej.')
        materialkoncowek.paragraph_format.left_indent = (Cm(4))
        materialkoncowek.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif zczegok2 == 2 and zczegot2 == 1:
        materialkoncowek = document.add_paragraph('Końcówki zostały wykonane ze stali nierdzewnej. Tuleje zostały wykonane ze stali węglowej.')
        materialkoncowek.paragraph_format.left_indent = (Cm(4))
        materialkoncowek.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif zczegok2 == 1 and zczegot2 == 2:
        materialkoncowek = document.add_paragraph('Końcówki zostały wykonane ze stali węglowej. Tuleje zostały wykonane ze stali nierdzewnej.')
        materialkoncowek.paragraph_format.left_indent = (Cm(4))
        materialkoncowek.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        print('errror')

    if uwagi == '':
        pass
    else:
        wpisaneuwagi = document.add_paragraph(uwagi)
        wpisaneuwagi.paragraph_format.left_indent = (Cm(4))
        wpisaneuwagi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_paragraph()

    opisweza = document.add_paragraph(kol3)
    opisweza.paragraph_format.left_indent = (Cm(4))
    opisweza.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_paragraph()

    if kol6 == '':
        pass
    else:
        wwew = document.add_paragraph('Warstwa wewn.:\t'+kol6)
        wwew.paragraph_format.left_indent = (Cm(4))

    if kol5 == '':
        pass
    else:
        wzmoc = document.add_paragraph('Wzmocnienie:\t'+kol5)
        wzmoc.paragraph_format.left_indent = (Cm(4))

    if kol4 == '':
        pass
    else:
        wzew = document.add_paragraph('Warstwa zewn.:\t'+kol4)
        wzew.paragraph_format.left_indent = (Cm(4))

    document.add_paragraph()

    if rodzdlu == 'Długość przewodu (mm)':
        ilo = document.add_paragraph('Ilość:\t' + str(ile) + ' ' + checkIle + ' o długości ' + dlu + ' mm. Tolerancja długości '+toler+'.')
        ilo.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()
    else:
        ilo = document.add_paragraph('Ilość:\t' + str(ile) + ' ' + checkIle + ' o długości węża ' + dlu + ' mm. Tolerancja długości '+toler+'.')
        ilo.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()

    if kol13 != '':
        document.add_paragraph('Katalogowe parametry pracy węża:')

        parametryweza0 = document.add_paragraph('\twarunki statyczne / dynamiczne')
        parametryweza0.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryweza0.paragraph_format.left_indent = (Cm(4))

        parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
        parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryweza1.paragraph_format.left_indent = (Cm(4))

        parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t' + str(kol11) + ' bar')
        parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryweza2.paragraph_format.left_indent = (Cm(4))

        parametryweza3 = document.add_paragraph('- promień zagięcia:\t' + str(kol12) + ' mm / '+str(kol13)+' mm')
        parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryweza3.paragraph_format.left_indent = (Cm(4))

        parametryweza4 = document.add_paragraph('- temperatura pracy:\tod ' + str(int(kol8)) + '°C do +' + str(int(kol7)) + '°C')
        parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryweza4.paragraph_format.left_indent = (Cm(4))

    else:
        if kol1 == ' ': #pusty indeks
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t ')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- promień zagięcia:\t')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph('- temperatura pracy:\t')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        elif kol11 == '': #brak w katalogu ciśnienia rozerwania
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- promień zagięcia:\t' + str(kol12) + ' mm')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperatura pracy:\tod ' + str(kol8) + '°C do +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        elif kol12 == '': #brak w katalogu promienia zagięcia
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t' + str(kol11) + ' bar')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperatura pracy:\tod ' + str(kol8) + '°C do +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        else: #wszystko jest
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t' + str(kol11) + ' bar')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- promień zagięcia:\t' + str(kol12) + ' mm')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperatura pracy:\tod ' + str(kol8) + '°C do +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
            parametryweza4.paragraph_format.left_indent = (Cm(4))


    document.add_paragraph()
    document.add_paragraph('Parametry pracy określone przez zamawiającego:')

    parametryklienta1 = document.add_paragraph('- ciśnienie robocze:\t'+cisR+' bar')
    parametryklienta1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
    parametryklienta1.paragraph_format.left_indent = (Cm(4))

    parametryklienta2 = document.add_paragraph('- medium:\t'+med)
    parametryklienta2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
    parametryklienta2.paragraph_format.left_indent = (Cm(4))

    if tw == 'otoczenia':
        parametryklienta3 = document.add_paragraph('- temp. wewnętrzna:\t' + tw)
        parametryklienta3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryklienta3.paragraph_format.left_indent = (Cm(4))
    else:
        parametryklienta3 = document.add_paragraph('- temp. wewnętrzna:\t' + tw + '°C')
        parametryklienta3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryklienta3.paragraph_format.left_indent = (Cm(4))

    if tz == 'otoczenia':
        parametryklienta4 = document.add_paragraph('- temp. zewnętrzna:\t'+tz)
        parametryklienta4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryklienta4.paragraph_format.left_indent = (Cm(4))
    else:
        parametryklienta4 = document.add_paragraph('- temp. zewnętrzna:\t' + tz + '°C')
        parametryklienta4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER)
        parametryklienta4.paragraph_format.left_indent = (Cm(4))

    document.add_paragraph()

    czytest2 = czytest.get()
    if czytest2 == 1:
        if cz == '1':
            ktoraminuta = 'minutę'
        elif cz == '2' or cz == '3' or cz == '4':
            ktoraminuta = 'minuty'
        else:
            ktoraminuta = 'minut'
        teksttestu1 = document.add_paragraph()
        teksttestu1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        teksttestu1.add_run(checkIle2+' testowano ' + medB + ', pod ciśnieniem ' + cisT + ' bar przez ' + cz + ' '+ktoraminuta+'.').font.underline = TRUE

        teksttestu2 = document.add_paragraph()
        teksttestu2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        teksttestu2.add_run('Nieszczelności nie stwierdzono.').font.underline = TRUE
        document.add_paragraph()
    else:
        pass

    if czyodtlu.get() == 1:
        tekstodtluszczania = document.add_paragraph()
        tekstodtluszczania.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tekstodtluszczania.add_run(checkIle2+' poddano procesowi wstępnego odtłuszczania do pracy z tlenem.').font.underline = TRUE
        tekstodtluszczania2 = document.add_paragraph()
        tekstodtluszczania2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tekstodtluszczania2.add_run('Odtłuszczenie finalne i dopuszczenie do eksploatacji leży po stronie zamawiającego.').font.underline = TRUE
        document.add_paragraph()
    else:
        pass

    if czyzuzycie.get() == 1:
        if ile == 1:
            tekstzuzycia = document.add_paragraph()
            tekstzuzycia.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tekstzuzycia.add_run('Ze względu na różnorodność mediów oraz ich agresywność, żywotność przewodu może być ograniczona.').font.underline = TRUE
            document.add_paragraph()
        else:
            tekstzuzycia = document.add_paragraph()
            tekstzuzycia.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tekstzuzycia.add_run('Ze względu na różnorodność mediów oraz ich agresywność, żywotność przewodów może być ograniczona.').font.underline = TRUE
            document.add_paragraph()
    else:
        pass

    if czyciaglosc.get() == 1:
        if ile == 1:
            tekstciaglosc = document.add_paragraph()
            tekstciaglosc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tekstciaglosc.add_run(
                'Potwierdzono badaniem zachowanie ciągłości elektrycznej pomiędzy końcówkami przewodu.').font.underline = TRUE
            document.add_paragraph()
        else:
            tekstciaglosc = document.add_paragraph()
            tekstciaglosc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tekstciaglosc.add_run(
                'Potwierdzono badaniem zachowanie ciągłości elektrycznej pomiędzy końcówkami przewodów.').font.underline = TRUE
            document.add_paragraph()
    else:
        pass

    if ile == 1:
        document.add_paragraph('Przewód może być użytkowany w parametrach określonych przez zamawiającego.')
    else:
        document.add_paragraph('Przewody mogą być użytkowane w parametrach określonych przez zamawiającego.')
    document.add_paragraph()

    if ile == 1:
        uzytkowanie1 = document.add_paragraph('Użytkowanie przewodu wg stanu technicznego - biorąc pod uwagę mechaniczne zużycie, ewentualne przetarcia podczas użytkowania, stopień ryzyka dla obsługi i wszystkie inne czynniki mogące mieć wpływ na zastosowanie i żywotność przewodów.')
        uzytkowanie1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        uzytkowanie1 = document.add_paragraph('Użytkowanie przewodów wg stanu technicznego - biorąc pod uwagę mechaniczne zużycie, ewentualne przetarcia podczas użytkowania, stopień ryzyka dla obsługi i wszystkie inne czynniki mogące mieć wpływ na zastosowanie i żywotność przewodów.')
        uzytkowanie1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    document.add_paragraph()

    przygzatw = document.add_paragraph('Przygotował:\t' + przygotowal + '\tZatwierdził:\t' + zatwierdzil)
    przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))
    przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
    przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(13.75))

    datydaty = document.add_paragraph('Data:\t' + kiedy_p + '\tData:\t' + kiedy_z)
    datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))
    datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
    datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(13.75))

    podpispodpis = document.add_paragraph('Podpis:\tPodpis:')
    podpispodpis.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
    podpispodpis.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))

    def filesave():
        filename = filedialog.asksaveasfilename(defaultextension='.docx', initialfile=nQC)
        if not filename: return
        document.save(filename)
        messagebox.showinfo("Dokument zapisano!", "Gratuluję sukcesu!")
    filesave()

menu = Menu(root)

ramkaKlient = LabelFrame(root, text="Informacje o kliencie")
ramkaKlient.columnconfigure(1, minsize=160)
ramkaKlient.grid(padx=5, pady=2, ipadx=5, ipady=5, sticky=N+S+E+W, column=0, row=0)

labelKlient1 = Label(ramkaKlient, text="Nazwa klienta:")
labelKlient1.grid(padx=2, pady=2, sticky=W, row=0)

entryKlient1 = Entry(ramkaKlient, width=33)
entryKlient1.grid(padx=2, pady=2, row=0, column=1, sticky=E)
entryKlient1.focus()
entryKlient1.bind('<Return>', znajdzklientaenter)

labelKlient2 = Label(ramkaKlient, text="Adres 1:")
labelKlient2.grid(padx=2, pady=2, sticky=W, row=1)

entryKlient2 = Entry(ramkaKlient, width=33)
entryKlient2.grid(padx=2, pady=2, row=1, column=1, sticky=E)

labelKlient3 = Label(ramkaKlient, text="Adres 2:")
labelKlient3.grid(padx=2, pady=2, sticky=W, row=2)

entryKlient3 = Entry(ramkaKlient, width=33)
entryKlient3.grid(padx=2, pady=2, row=2, column=1, sticky=E)

Button(ramkaKlient, text='Znajdź klienta', command=znajdzklienta).grid(padx=2, pady=2, row=3, column=0, columnspan=2, sticky=W)
Button(ramkaKlient, text="Dodaj klienta", command=dodajklienta).grid(padx=2, pady=2, row=3, column=0, columnspan=2, sticky=E)

gornaRamka = LabelFrame(root, text="Ważne numery")
gornaRamka.columnconfigure(1, minsize=160)
gornaRamka.grid(padx=5, pady=2, ipadx=5, ipady=3, sticky=N + S + E + W, column=0, row=1)

tekstnumerQC = Label(gornaRamka, text="Numer świadectwa QC")
tekstnumerQC.grid(padx=2, pady=2, sticky=W, row=0)

numerQC = Entry(gornaRamka)
numerQC.insert(END, '1234')
numerQC.grid(padx=2, pady=2, row=0, column=1, sticky=E)

tekstnumerKRP = Label(gornaRamka, text="Numer KRP")
tekstnumerKRP.grid(padx=2, pady=2, sticky=W, row=1)

numerKRP1 = Entry(gornaRamka, width=3)
numerKRP1.insert(END, '50')
numerKRP1.grid(ipadx=5, padx=2, pady=2, row=1, column=1, sticky=W)

numerKRP2 = Entry(gornaRamka, width=6)
numerKRP2.insert(END, '4321')
numerKRP2.grid(ipadx=5, padx=2, pady=2, row=1, column=1)

numerKRP3 = Entry(gornaRamka, width=3)
numerKRP3.insert(END, time.strftime('%y'))
numerKRP3.grid(ipadx=5, padx=2, pady=2, row=1, column=1, sticky=E)

tekstnumerTU = Label(gornaRamka, text="Numer TU")
tekstnumerTU.grid(padx=2, pady=2, sticky=W, row=2)

tu = Entry(gornaRamka)
tu.insert(END, '50/16/1234')
tu.grid(padx=2, pady=2, row=2, column=1, sticky=E)

labelZamowienie = Label(gornaRamka, text='Zamówienie')
labelZamowienie.grid(padx=2, pady=2, sticky=W, row=3)

numerZam = Entry(gornaRamka, width=20)
numerZam.grid(padx=2, pady=2, column=1, row=3, sticky=E, columnspan=3)

koncowkaframe = LabelFrame(root, text='Końcówki przewodu')
koncowkaframe.grid(padx=5, pady=2, row=2, column=0, sticky=W+E+S+N)

koncowka1label = Label(koncowkaframe, text='Końcówka 1:')
koncowka1label.grid(padx=2, pady=2, sticky=W, column=0, row=1)

koncowka1entry = Entry(koncowkaframe, width=33)
koncowka1entry.grid(padx=2, pady=2, sticky=E, column=1, row=1)

koncowka2label = Label(koncowkaframe, text='Końcówka 2:')
koncowka2label.grid(padx=2, pady=2, sticky=W, column=0, row=2)

koncowka2entry = Entry(koncowkaframe, width=33)
koncowka2entry.grid(padx=2, pady=2, sticky=E, column=1, row=2)

def koncowkacheck():
    if ilekoncowek.get() == 2:
        koncowka2entry.config(state='disabled')
    elif ilekoncowek.get() == 1:
        koncowka2entry.config(state='normal')
    else:
        pass

ilekoncowek = IntVar()
dwiekoncowki = Radiobutton(koncowkaframe, text="Dwie różne", variable=ilekoncowek, value=1, command=koncowkacheck)
dwiekoncowki.grid(sticky=W, row=0, column=0)
dwiekoncowki.invoke()

jednakoncowka = Radiobutton(koncowkaframe, text="Takie same", variable=ilekoncowek, value=2, command=koncowkacheck)
jednakoncowka.grid(sticky=E, row=0, column=1)

materialramka = LabelFrame(root, text="Materiał końcówek i tulei")
materialramka.grid(padx=5, pady=2, ipadx=5, ipady=5, row=3, column=0, sticky=W+E+N+S)

materialklabel = Label(materialramka, text='Wybierz materiał końcówek:')
materialklabel.grid(padx=2, pady=2, sticky=W, column=0, row=0)

zczegok = IntVar()
koncowkaweglowa = Radiobutton(materialramka, text="Węglowa", variable=zczegok, value=1)
koncowkaweglowa.grid(sticky=W, row=1, column=0)
koncowkaweglowa.invoke()

koncowkanierdzewna = Radiobutton(materialramka, text="Nierdzewna", variable=zczegok, value=2)
koncowkanierdzewna.grid(sticky=E, row=1, column=1)

materialtlabel = Label(materialramka, text='Wybierz materiał tulei:')
materialtlabel.grid(padx=2, pady=2, sticky=W, column=0, row=2)

zczegot = IntVar()
tulejaweglowa = Radiobutton(materialramka, text="Węglowa", variable=zczegot, value=1)
tulejaweglowa.grid(sticky=W, row=3, column=0)
tulejaweglowa.invoke()

tulejanierdzewna = Radiobutton(materialramka, text="Nierdzewna", variable=zczegot, value=2)
tulejanierdzewna.grid(sticky=E, row=3, column=1)



dolnaRamka = LabelFrame(root, text="Informacje o przewodzie")
dolnaRamka.columnconfigure(0, minsize=170)
dolnaRamka.grid(ipadx=2, padx=5, pady=2, row=0, column=1, rowspan=3,sticky=N+S+E+W)

waz = ttk.Combobox(dolnaRamka, state='readonly', width=22)
waz['values'] = ('Długość przewodu (mm)', 'Długość węża (mm)')
waz.current([0])
waz.grid(padx=2, pady=2, row=2, column=0, sticky=W)

labelIlosc = Label(dolnaRamka, text="Ilość")
labelIlosc.grid(padx=2, pady=2, sticky=W, row=1)

ilosc = Entry(dolnaRamka)
ilosc.insert(END, '15')
ilosc.grid(padx=2, pady=2, row=1, column=0, sticky=E)

dlugosc = Entry(dolnaRamka)
dlugosc.grid(padx=2, pady=2, sticky=E, column=0, row=2)
dlugosc.insert(END, '100')

ramkawyboru = LabelFrame(dolnaRamka, text='Wybór węża')
ramkawyboru.grid(padx=5, pady=2, ipadx=5, row=3, column=0, sticky=N)

indeksweza = Entry(ramkawyboru, state='disabled', width=25)
indeksweza.grid(padx=2, pady=2, sticky=E, column=0, row=4)

labelwprowadzwaz = Label(ramkawyboru, text='Wprowadź indeks:\n(potwierdź Enterem)')
labelwprowadzwaz.grid(padx=2, pady=2, sticky=W, column=0, row=1)

szukajweza = Entry(ramkawyboru, width=23)
szukajweza.grid(padx=2, pady=2, sticky=E, column=0, row=1)
szukajweza.bind('<Return>', szukaj)

rezultaty = Listbox(ramkawyboru, height=11, activestyle='none', width=44)
rezultaty.grid(row=2, column=0, pady=2)
rezultaty.bind("<Double-Button-1>", kliknijwybor)

yscrollbar = Scrollbar(ramkawyboru, orient=VERTICAL)
yscrollbar.config(command=rezultaty.yview)
yscrollbar.grid(row=2, column=0, sticky=N+S+E, pady=2)

rezultaty.config(yscrollcommand=yscrollbar.set)

guziczek = Button(ramkawyboru, text='Wybierz wąż', command=wprowadzwybor)
guziczek.grid(ipadx=5,ipady=5, padx=10, pady=8, sticky=W, column=0, row=4)

warunkiKlienta = LabelFrame(root, text="Warunki robocze klienta")
warunkiKlienta.grid(padx=5, pady=2, row=0, column=2, sticky=E+W+N+S)

labelMedium = Label(warunkiKlienta, text="Medium")
labelMedium.grid(padx=2, pady=2, sticky=W)

medium = Entry(warunkiKlienta)
medium.grid(padx=2, pady=2, sticky=E, column=1, row=0)

labelCisRob = Label(warunkiKlienta, text="Ciśnienie robocze (bar)")
labelCisRob.grid(padx=2, pady=2, sticky=W, row=1)

cisnienieRob = Entry(warunkiKlienta)
cisnienieRob.grid(padx=2, pady=2, sticky=E, column=1, row=1)

labelTempWew = Label(warunkiKlienta, text="Temp. wewnętrzna (°C)")
labelTempWew.grid(padx=2, pady=2, sticky=W, row=2)

tempWew = Entry(warunkiKlienta)
tempWew.grid(padx=2, pady=2, sticky=E, column=1, row=2)

labelTempZew = Label(warunkiKlienta, text="Temp. zewnętrzna (°C)")
labelTempZew.grid(padx=2, pady=2, sticky=W, row=3)

tempZew = Entry(warunkiKlienta)
tempZew.grid(padx=2, pady=2, sticky=E, column=1, row=3)

warunkiTestu = LabelFrame(root, text="Warunki testu")
warunkiTestu.grid(padx=5, pady=2, sticky=W + E + N+S, row=1, column=2)

czymBadane = Label(warunkiTestu, text="Rodzaj testu")
czymBadane.grid(padx=2, pady=2, sticky=W, row=1, column=0)

mediumBadania = ttk.Combobox(warunkiTestu, state='readonly')
mediumBadania['value'] = ('hydrostatycznie', 'powietrzem pod wodą', 'azotem pod wodą')
mediumBadania.current(0)
mediumBadania.grid(padx=2, pady=2, row=1, column=1)

labelCisTest = Label(warunkiTestu, text="Ciśnienie testu (bar)")
labelCisTest.grid(padx=2, pady=2, sticky=W, row=3, column=0)

cisnienieTest = Entry(warunkiTestu)
cisnienieTest.grid(padx=2, pady=2, sticky=E, column=1, row=3)

labelCzasTestu = Label(warunkiTestu, text="Czas testu (min)")
labelCzasTestu.grid(padx=2, pady=2, sticky=W, row=4, column=0)

czasTestu = Entry(warunkiTestu)
czasTestu.grid(padx=2, pady=2, sticky=E, column=1, row=4)

def naccheck():
    if czytest.get() == 2:
        mediumBadania.config(state='disabled')
        cisnienieTest.config(state='disabled')
        czasTestu.config(state='disabled')
    elif czytest.get() == 1:
        mediumBadania.config(state='readonly')
        cisnienieTest.config(state='normal')
        czasTestu.config(state='normal')
    else:
        pass

czytest = IntVar()
badanieTak = Radiobutton(warunkiTestu, text="TAK", variable=czytest, value=1, command=naccheck)
badanieTak.grid(sticky=W, row=0, column=0)
badanieTak.invoke()

badanieNie = Radiobutton(warunkiTestu, text="NIE", variable=czytest, value=2, command=naccheck)
badanieNie.grid(sticky=W, row=0, column=1)


ktokiedyRamka = LabelFrame(root, text="Przygotował i zatwierdził")
ktokiedyRamka.columnconfigure(1, minsize=200)
ktokiedyRamka.grid(padx=5, pady=2, row=2, column=2)

labelprzygotowal = Label(ktokiedyRamka, text='Przygotował:')
labelprzygotowal.grid(padx=2, pady=2, sticky=W, row=0, column=0)

entryprzygotowal = ttk.Combobox(ktokiedyRamka, state='readonly')
entryprzygotowal['values'] = ['Bartłomiej Gątarski', 'Michał Rosada', 'Robert Tomaszewski', 'Mariusz Węcławiak', 'Krzysztof Wenda', 'Hubert Stanisławski', 'Daniel Józefiak']
entryprzygotowal.grid(padx=2, pady=2, sticky=E, column=1, row=0)
entryprzygotowal.current(0)

labelkiedyprzygotowal = Label(ktokiedyRamka, text='Data:')
labelkiedyprzygotowal.grid(padx=2, pady=2, sticky=W, column=0, row=1)

now = datetime.datetime.now()
biezacadata = str(now.day)+'.'+str(now.month)+'.'+str(now.year)

entrykiedyprzygotowal = Entry(ktokiedyRamka)
entrykiedyprzygotowal.insert(END, biezacadata)
entrykiedyprzygotowal.grid(padx=2, pady=2, sticky=E, column=1, row=1)

labelzatwierdzil = Label(ktokiedyRamka, text='Zatwierdził:')
labelzatwierdzil.grid(padx=2, pady=2, sticky=W, row=2, column=0)

entryzatwierdził = ttk.Combobox(ktokiedyRamka, state='readonly')
entryzatwierdził['values'] = ['Hubert Stanisławski', 'Krzysztof Wenda', 'Daniel Józefiak']
entryzatwierdził.grid(padx=2, pady=2, sticky=E, column=1, row=2)
entryzatwierdził.current(0)

labelkiedyzatwierdzil = Label(ktokiedyRamka, text='Data:')
labelkiedyzatwierdzil.grid(padx=2, pady=2, sticky=W, column=0, row=3)

entrykiedyzatwierdzil = Entry(ktokiedyRamka)
entrykiedyzatwierdzil.insert(END, biezacadata)
entrykiedyzatwierdzil.grid(padx=2, pady=2, sticky=E, column=1, row=3)


ramkauwagi = LabelFrame(root, text='Uwagi / adnotacje (pod informacją o końcówkach):')
ramkauwagi.grid(padx=5, pady=2, row=3, column=1, sticky=W+E+S+N)

poleuwagi = Text(ramkauwagi, height=6, width=33)
poleuwagi.grid(padx=9, pady=5)

scrolluwagi = Scrollbar(ramkauwagi, orient=VERTICAL)
scrolluwagi.config(command=poleuwagi.yview)
scrolluwagi.grid(row=0, column=0, sticky=N+S+E, pady=5)

poleuwagi.config(yscrollcommand=scrolluwagi.set)


ramkaoznaczenia = LabelFrame(root, text='Niestandardowe oznaczenie')
ramkaoznaczenia.grid(padx=5, pady=2, row=4, column=1, sticky=E+W)

typoznaczenia = ttk.Combobox(ramkaoznaczenia, state='readonly', width=45)
typoznaczenia['values'] = ['NIE - oznaczenie standardowe', 'Pharmaline N/G', 'CE I (stalowy)', 'CE II (stalowy)', 'Spir Star', 'CE I (niestalowy)']
typoznaczenia.grid(padx=4, pady=2, column=0, row=0)
typoznaczenia.bind('<<ComboboxSelected>>', onselectoznakowanie)
typoznaczenia.current(0)

poleoznaczenia = Text(ramkaoznaczenia, width=35, height=4)
poleoznaczenia.grid(padx=4, pady=5, column=0, row=1)


dodatkoweinformacje = LabelFrame(root, text='Dodatkowe informacje')
dodatkoweinformacje.columnconfigure(0, minsize=270)
dodatkoweinformacje.grid(padx=5, pady=2, row=3, column=2, sticky=E+W)

czyodtlu = IntVar()
odtluszczane = Checkbutton(dodatkoweinformacje, text='Notka o odtłuszczaniu', variable=czyodtlu).grid(padx=2, pady=2, sticky=W, column=0, row=0)

czyzuzycie = IntVar()
ograniczonezycie = Checkbutton(dodatkoweinformacje, text='Notatka o ograniczonej żywotności', variable=czyzuzycie).grid(padx=2, pady=2, sticky=W, column=0, row=1)

czyciaglosc = IntVar()
ciagloscelektryczna = Checkbutton(dodatkoweinformacje, text='Notatka o ciągłości elektrycznej przewodu', variable=czyciaglosc).grid(padx=2, pady=2, sticky=W, column=0, row=2)

przyciskisterowania = LabelFrame(root)
przyciskisterowania.grid(padx=5, pady=2, row=4, column=2, sticky=E+W+N+S)

przyciskGeneruj = Button(przyciskisterowania, text="ZAPISZ\nŚWIADECTWO\n(wersja PL)", height=4, width=15, command=generujpopolsku)
przyciskGeneruj.grid(padx=30, pady=20, column=1, row=0)

przyciskwyczysc = Button(przyciskisterowania, text="Wyczyść", height=2, width=10, command=wyczyscpola)
przyciskwyczysc.grid(padx=10, pady=30, column=0, row=0, sticky=N+S)

dodajnowywaz = Button(root, text='Dodaj nowy wąż', command=otworznowywaz)
dodajnowywaz.grid(row=4, column=0)

root.mainloop()