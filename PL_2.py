from cryptography.fernet import Fernet
from datetime import datetime
from time import strftime
from tkinter import Entry, Label, LabelFrame, Tk, Listbox, N, S, E, W, NSEW, Button, Radiobutton, Checkbutton, END, ANCHOR, LEFT, Scrollbar, Toplevel, VERTICAL, IntVar, StringVar, Text, TRUE, RAISED, FLAT, Frame, SOLID, font
from tkinter import Menu
from tkinter import messagebox
from tkinter import ttk
from tkinter import filedialog
from docx import Document
from xlutils.copy import copy
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.shared import Cm
import xlrd
from os import startfile, getlogin
import pyperclip
from PIL.ImageTk import PhotoImage, Image
import webbrowser
from distutils.version import LooseVersion
from configparser import ConfigParser


def patch_crypto_be_discovery():

    """
    Monkey patches cryptography's backend detection.
    Objective: support pyinstaller freezing.
    """

    from cryptography.hazmat import backends

    try:
        from cryptography.hazmat.backends.commoncrypto.backend import backend as be_cc
    except ImportError:
        be_cc = None
    try:
        from cryptography.hazmat.backends.openssl.backend import backend as be_ossl
    except ImportError:
        be_ossl = None

    backends._available_backends_list = [
        be for be in (be_cc, be_ossl) if be is not None
    ]

patch_crypto_be_discovery()

root = Tk()

root.wm_minsize(923, 0)
root.iconbitmap(r'bin\tubes.ico')

versionnr = LooseVersion('0.5.9')

Config = ConfigParser()

def ConfigSectionMap(section):
    Config.read('config.ini')
    dict1 = {}
    options = Config.options(section)
    for option in options:
        try:
            dict1[option] = Config.get(section, option)
        except:
            print("exception on %s!" % option)
            dict1[option] = None
    return dict1

def downloadupdate():
    startfile('aktualizator.exe')
    root.destroy()


def checkupdate():
    A = open('bin/adres.txt').read()
    try:
        x = open(A+r'/version.tub')
        y = x.readline()
        if 'version' in y:
            _, z = y.split('=')
        z1 = LooseVersion(z)
        if z1 > versionnr:
            p1 = messagebox.askyesno("Aktualizacja", "Znaleziono nową wersję programu.\nCzy chcesz ją teraz pobrać?", default="yes")
            if p1:
                downloadupdate()
        else:
            pass
    except FileNotFoundError:
        pass

def wczytajoddzial():
    try:
        c = ConfigSectionMap('Uzytkownik')['oddzial']
        ktoryoddzial.set(c)
        d = ConfigSectionMap('Uzytkownik')['tooltip']
        wlaczniktooltip.set(d)
        e = ConfigSectionMap('Uzytkownik')['test']
        wlaczniktest.set(e)
    except:
        ktoryoddzial.set('50')
        Config.read('config.ini')
        Config.add_section('Uzytkownik')
        Config.set('Uzytkownik', 'oddzial', '50')
        Config.set('Uzytkownik', 'tooltip', '0')
        Config.set('Uzytkownik', 'test', '1')
        x = open(r'config.ini', 'w')
        Config.write(x)
        messagebox.showwarning('Nie znaleziono konfiguracji oddziału!','Nie znaleziono konfiguracji oddziału.\nPrzywrócono domyślną konfigurację - PL50.\nAby ustawić inny oddział, skorzystaj z menu.')

menubar = Menu(root)
ktoryoddzial = StringVar()
wlaczniktooltip = IntVar()
wlaczniktest = IntVar()

def setnumber():
    try:
        Config.read('config.ini')
        Config.set('Uzytkownik', 'oddzial', str(ktoryoddzial.get()))
        x = open('config.ini', 'w')
        Config.write(x)
        root.title("Świadectwomat v"+str(versionnr)+" - "+slownikoddzialow.get(ktoryoddzial.get()))
        entryprzygotowal['values'] = listaprzygotowal.get(ktoryoddzial.get())
        entryprzygotowal.current(0)
        kimjestes()
        setzatwierdzil(ktoryoddzial.get())
    except:
        pass


def settooltip():
    try:
        Config.read('config.ini')
        Config.set('Uzytkownik', 'tooltip', str(wlaczniktooltip.get()))
        x = open('config.ini', 'w')
        Config.write(x)
    except:
        pass

def setdefaulttest():
    try:
        Config.read('config.ini')
        Config.set('Uzytkownik', 'test', str(wlaczniktest.get()))
        x = open('config.ini', 'w')
        Config.write(x)
    except:
        pass

glownemenu = Menu(menubar, tearoff=0)

oddzialmenu = Menu(menubar,tearoff=0)
# oddzialmenu.add_radiobutton(label='02 - Kalisz', variable=ktoryoddzial, value='02', command=setnumber)
oddzialmenu.add_radiobutton(label='04 - Katowice', variable=ktoryoddzial, value='04', command=setnumber)
oddzialmenu.add_radiobutton(label='05 - Gdańsk', variable=ktoryoddzial, value='05', command=setnumber)
oddzialmenu.add_radiobutton(label='07 - Poznań', variable=ktoryoddzial, value='07', command=setnumber)
oddzialmenu.add_radiobutton(label='15 - Toruń', variable=ktoryoddzial, value='15', command=setnumber)
oddzialmenu.add_radiobutton(label='18 - Częstochowa', variable=ktoryoddzial, value='18', command=setnumber)
oddzialmenu.add_radiobutton(label='20 - Olsztyn', variable=ktoryoddzial, value='20', command=setnumber)
oddzialmenu.add_radiobutton(label='50 - Centrala', variable=ktoryoddzial, value='50', command=setnumber)
oddzialmenu.add_radiobutton(label='61 - Produkcja OEM', variable=ktoryoddzial, value='61', command=setnumber)
menubar.add_cascade(label="Plik", menu=glownemenu)
glownemenu.add_cascade(label='Oddział', menu=oddzialmenu)
glownemenu.add_separator()

glownemenu.add_checkbutton(label='Podpowiedzi', variable=wlaczniktooltip, command=settooltip)
glownemenu.add_separator()
glownemenu.add_checkbutton(label='Domyślnie włączone testowanie ciśń.', variable=wlaczniktest, command=setdefaulttest)


glownemenu.add_separator()
glownemenu.add_command(label="Wyjście", command=root.quit)

root.config(menu=menubar)

listawezy = []
listaparametrow = ['0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0','0']

listaklientow = []

def ladujklientow():
    try:
        wbX = xlrd.open_workbook(
            '//Srv-dokumenty/Swiadectwomat/bin/klient.xls',
            formatting_info=True)
    except FileNotFoundError:
        wbX = xlrd.open_workbook('bin/klient.xls', formatting_info=True)
        messagebox.showwarning('Twój INTERNET mówi NIE!',
                               'Nie można nawiązać połączenia ze spisem klientów na serwerze sieciowym!\n\nZaładowano listę klientów z zapasowego pliku lokalnego.\n\nUWAGA - funkcja dodania nowego klienta do listy jest WYŁĄCZONA!')
    global wsX
    wsX = wbX.sheet_by_index(0)
    listaklientow.clear()
    for row in wsX.col(0):
        key = b'Uf-FryHfa-JXDToRDClFsSrv2h56kzEUCK5imvQngbo='
        f = Fernet(key)
        row_value = row.value
        kodowanie = row_value.encode()
        odkryte = f.decrypt(kodowanie)
        odszyfrowane = odkryte.decode('utf-8')
        listaklientow.append(odszyfrowane)

ladujklientow()

def dodajklienta():
    try:
        wbnX = xlrd.open_workbook(
            '//Srv-dokumenty/Swiadectwomat/bin/klient.xls',
            formatting_info=True)
        adresklienta = '//Srv-dokumenty/Swiadectwomat/bin/klient.xls'
    except FileNotFoundError:
        messagebox.showerror('Zapis lokalny niemożliwy!',
                               'Dodanie klienta do bazy nie jest możliwe ze względu na brak połączenia ze spisem na serwerze sieciowym!')
        pass
    wsnX = wbnX.sheet_by_index(0)
    wbn1X = copy(wbnX)
    wsn1X = wbn1X.get_sheet(0)
    liczbawierszy = 0
    for _ in wsnX.col(0):
        liczbawierszy += 1
    dk1 = entryKlient1.get()
    dk2 = entryKlient2.get()
    dk3 = entryKlient3.get()
    key = b'Uf-FryHfa-JXDToRDClFsSrv2h56kzEUCK5imvQngbo='
    f = Fernet(key)
    dk1token = f.encrypt(bytes(dk1, 'utf-8'))
    dk2token = f.encrypt(bytes(dk2, 'utf-8'))
    dk3token = f.encrypt(bytes(dk3, 'utf-8'))
    dk11token = dk1token.decode('utf-8')
    dk22token = dk2token.decode('utf-8')
    dk33token = dk3token.decode('utf-8')
    daneklienta = [dk11token, dk22token, dk33token]
    potwierdzeniedodaniaklienta = messagebox.askyesno('Dodanie nowego klienta', 'Czy na pewno chcesz dodać następującego klienta:\n\n'+dk1+'\n'+dk2+'\n'+dk3)
    if potwierdzeniedodaniaklienta:
        for col_index in range(0,3):
            wsn1X.write(liczbawierszy, col_index, daneklienta[col_index])
        wbn1X.save(adresklienta)
        messagebox.showinfo("Dodano klienta do listy!", "Dodano klienta do listy!\n\nGratuluję sukcesu!")
        ladujklientow()
    else:
        pass

tree_columns = ("Nazwa klienta", "Adres 1", "Adres 2")
def sortby(tree, col, descending):
    """Sort tree contents when a column is clicked on."""
    # grab values to sort
    data = [(tree.set(child, col), child) for child in tree.get_children('')]

    # reorder data
    data.sort(reverse=descending)
    for indx, item in enumerate(data):
        tree.move(item[1], '', indx)

    # switch the heading so that it will sort in the opposite direction
    tree.heading(col,
        command=lambda col=col: sortby(tree, col, int(not descending)))

class Drzewo(object):
    def __init__(self, rodzic, tree_data, entry1, entry2, entry3):
        self.tree = None
        self._setup_widgets(rodzic)
        self._build_tree(tree_data)
        self.tree.bind('<Double-1>', self.daneklienta)

    def _setup_widgets(self, rodzic):
        container = ttk.Frame(rodzic)
#         container.pack(fill='both', expand=True)
        container.grid()

        # XXX Sounds like a good support class would be one for constructing
        #     a treeview with scrollbars.
        self.tree = ttk.Treeview(rodzic, columns=tree_columns, show="headings")
        vsb = ttk.Scrollbar(rodzic, orient="vertical", command=self.tree.yview)
        hsb = ttk.Scrollbar(rodzic, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self.tree.grid(column=0, row=0, sticky='nsew')
        vsb.grid(column=1, row=0, sticky='ns')
        hsb.grid(column=0, row=1, sticky='ew')

        container.grid_columnconfigure(0, weight=1)
        container.grid_rowconfigure(0, weight=1)

    def _build_tree(self, tree_data):
        for col in tree_columns:
            self.tree.heading(col, text=col.title(),
                command=lambda c=col: sortby(self.tree, c, 0))
            # XXX tkFont.Font().measure expected args are incorrect according
            #     to the Tk docs
            self.tree.column(col, width=font.Font().measure(col.title()))

        for item in tree_data:
            self.tree.insert('', 'end', values=item)

            # adjust columns lenghts if necessary
            for indx, val in enumerate(item):
                ilen = font.Font().measure(val)
                if self.tree.column(tree_columns[indx], width=None) < ilen:
                    self.tree.column(tree_columns[indx], width=ilen)
    
    def daneklienta(self, event):
        self.entryKlient1 = entryKlient1
        self.entryKlient2 = entryKlient2
        self.entryKlient3 = entryKlient3
        item = self.tree.identify_row(event.y)
        entryKlient1.delete(0, END)
        entryKlient1.insert(END, self.tree.item(item, "values")[0])
        entryKlient1.validate()
        entryKlient2.delete(0, END)
        entryKlient2.insert(END, self.tree.item(item, "values")[1])
        entryKlient2.validate()
        entryKlient3.delete(0, END)
        entryKlient3.insert(END, self.tree.item(item, "values")[2])
        entryKlient3.validate()
        oknoklient.destroy()
        ladujklientow()

    def bind(self, k, f):
        return self.tree.bind(k, f)

def znajdzklienta():

    wprowadzony = entryKlient1.get()
    listaponumerowana = list(enumerate(listaklientow))
    znalezioneindeksy = [i for i,x in listaponumerowana if wprowadzony.lower() in x.lower()]
    rezultatyklient = []
    if not znalezioneindeksy:
        messagebox.showinfo("Nie znaleziono klienta!", "Nie udało się dopasować wprowadzonej frazy do żadnego z klientów.\nWypełnij pola adresowe ręcznie, a następnie kliknij przycisk 'Dodaj klienta', aby zachować dane na przyszłość.")
    else:
        global oknoklient
        oknoklient = Toplevel(root)
        oknoklient.iconbitmap(r'bin\tubes.ico')
        oknoklient.title('Klienci zawierający szukane słowo:')
        for item in znalezioneindeksy:
            wartoscadresu1 = wsX.cell(item,1).value
            key = b'Uf-FryHfa-JXDToRDClFsSrv2h56kzEUCK5imvQngbo='
            f = Fernet(key)
            kodowanie1 = wartoscadresu1.encode()
            odkryte1 = f.decrypt(kodowanie1)
            wartoscadresu11 = odkryte1.decode('utf-8')
            wartoscadresu2 = wsX.cell(item,2).value
            kodowanie2 = wartoscadresu2.encode()
            odkryte2 = f.decrypt(kodowanie2)
            wartoscadresu22 = odkryte2.decode('utf-8')
            rezultatyklient.append((listaklientow[item], wartoscadresu11, wartoscadresu22))
        
        Drzewo(oknoklient, rezultatyklient, entryKlient1, entryKlient2, entryKlient3)
        Tooltip(oknoklient, text='Potwierdź wybór klienta podwójnym kliknięciem')
    
        if len(rezultatyklient) == 1:
            tupelek = rezultatyklient[0]
            entryKlient1.delete(0, END)
            entryKlient1.insert(END, tupelek[0])
            entryKlient1.validate()
            entryKlient2.delete(0, END)
            entryKlient2.insert(END, tupelek[1])
            entryKlient2.validate()
            entryKlient3.delete(0, END)
            entryKlient3.insert(END, tupelek[2])
            entryKlient3.validate()
            oknoklient.destroy()
            ladujklientow()

def znajdzklientaenter(event):
    znajdzklienta()

ludzie = {
    'adarog': 0,
    'damsre': 1,
    "arkkla": 2,
    "marwec": 3,
    "micros": 4,
    "robtom": 5,
    "krzwen": 6,
    "hubsta": 7,
    "danjoz": 8,
    "krzszr": 0,
    "malsmu": 1,
    "marfal": 2,
    "robwoz": 3,
    "agnbod": 0,
    "jarwal": 1,
    "tomrec": 2,
    "sylpik": 0,
    "wojlej": 1,
    "arkgaj": 2,
    "miczah": 3,
    "piosik": 4,
    "piomur": 0,
    "pawsko": 1,
    "martry": 2,
    "joasta": 3,
    "tomnap": 0,
    "eugpod": 1,
    "darken": 0,
    "przkap": 1,
    "tomgrz": 0,
    "jaknec": 1,
    "pawpud": 2
    }



def kimjestes():
    try:
        entryprzygotowal.current(ludzie.get(getlogin()))
    except:
        pass

def ladujlistewezy(s):
    if s == 'PL':
        try:
            wb = xlrd.open_workbook(
                '//Srv-dokumenty/Swiadectwomat/bin/spis.xls')
        except FileNotFoundError:
            wb = xlrd.open_workbook('bin/spis.xls')
            messagebox.showwarning('Nie można połączyć ze spisem węży',
                                   'Nie udało się pobrać listy węży z folderu sieciowego. Załadowano listę z dysku lokalnego, która może zawierać nieaktualne dane.\nW razie kolejnych niepowodzeń zgłoś błąd!')
        global ws
        ws = wb.sheet_by_name('lista')
        listawezy.clear()
        for row in ws.col(0):
            listawezy.append(row.value)
    elif s == 'ENG':
        try:
            wb = xlrd.open_workbook(
                '//Srv-dokumenty/Swiadectwomat/bin/spis_en.xls')
        except FileNotFoundError:
            wb = xlrd.open_workbook('bin/spis_en.xls')
            messagebox.showwarning('Nie można połączyć ze spisem węży',
                                   'Nie udało się pobrać listy węży z folderu sieciowego. Załadowano listę z dysku lokalnego, która może zawierać nieaktualne dane.\nW razie kolejnych niepowodzeń zgłoś błąd!')
        ws = wb.sheet_by_name('lista')
        listawezy.clear()
        for row in ws.col(0):
            listawezy.append(row.value)
    else:
        print('cos sie popsulo')

ladujlistewezy('PL')

def otworztoler(flaga):
    spistoler = {
        'P': 'toler-przewody',
        'S': 'toler-stalowe',
        'K': 'toler-kompozyt',
        'C': 'toler-corroflon'
        }
    startfile('bin\\'+spistoler[flaga]+'.pdf')

def otworznowywaz():
    matka = Toplevel()

    matka.title('Dodaj opis przewodu')
    matka.iconbitmap(r'bin\tubes.ico')

    def dodajwazdolisty():
        global wersjajezykowa
        if wersjajezykowa.get() == 'PL':
            try:
                wbn = xlrd.open_workbook('//Srv-dokumenty/Swiadectwomat/bin/spis.xls', formatting_info=True)
                adresspisu = '//Srv-dokumenty/Swiadectwomat/bin/spis.xls'
            except FileNotFoundError:
                wbn = xlrd.open_workbook('bin/spis.xls', formatting_info=True)
                messagebox.showwarning('Zapis lokalny!','Wąż dodany do listy lokalnej, zapasowej. Wprowadzone informacje o wężu NIE będą utrwalone w spisie sieciowym!')
                adresspisu = 'bin/spis.xls'
            wsn = wbn.sheet_by_name('lista')
            wbn1 = copy(wbn)
            wsn1 = wbn1.get_sheet(0)
        elif wersjajezykowa.get() == 'ENG':
            try:
                wbn = xlrd.open_workbook('//Srv-dokumenty/Swiadectwomat/bin/spis_en.xls', formatting_info=True)
                adresspisu = '//Srv-dokumenty/Swiadectwomat/bin/spis_en.xls'
            except FileNotFoundError:
                wbn = xlrd.open_workbook('bin/spis_en.xls', formatting_info=True)
                messagebox.showwarning('Zapis lokalny!','Wąż dodany do listy lokalnej, zapasowej. Wprowadzone informacje o wężu NIE będą utrwalone w spisie sieciowym!')
                adresspisu = 'bin/spis_en.xls'
            wsn = wbn.sheet_by_name('lista')
            wbn1 = copy(wbn)
            wsn1 = wbn1.get_sheet(0)
        liczbawierszy = 0
        for _ in wsn.col(0):
            liczbawierszy += 1
        k1 = kolumnaindeks.get()
        k2 = kolumnatypweza.get()
        k3 = kolumnaopisweza.get()
        k4 = kolumnawarstwazewnetrzna.get()
        k5 = kolumnawzmocnienie.get()
        k6 = kolumnawarstwawewnetrzna.get()
        k7 = kolumnatempmaksymalna.get()
        k8 = kolumnatempminimalna.get()
        k9 = kolumnasrednica.get()
        k10 = kolumnacisnienierobocze.get()
        k11 = kolumnacisnienierozerwania.get()
        k12 = kolumnapromienstatyczny.get()
        k13 = kolumnapromiendynamiczny.get()
        k14 = kolumnaflagatolerancji.get()
        parametrydododania = [k1, k2, k3, k4, k5, k6, k7, k8, k9, k10, k11, k12, k13, k14]
        for col_index in range(0,14):
            wsn1.write(liczbawierszy, col_index, parametrydododania[col_index])
        wbn1.save(adresspisu)
        messagebox.showinfo("Dodano wąż do listy!", "Dodano wąż do listy!\n\nGratuluję sukcesu!")
        ladujlistewezy(wersjajezykowa.get())
        matka.destroy()


    Label(matka, text='Indeks węża:').grid(padx=2, pady=2, sticky=W, row=0, column=0)

    kolumnaindeks = Entry(matka, width=30)
    kolumnaindeks.grid(padx=2, pady=2, sticky=E, row=0, column=1)

    Label(matka, text='Typ węża (np. 2SN, PARRAP, MP 20):').grid(padx=2, pady=2, sticky=W, row=1, column=0)

    kolumnatypweza = Entry(matka, width=30)
    kolumnatypweza.grid(padx=2, pady=2, sticky=E, row=1, column=1)

    Label(matka, text='Opis węża:').grid(padx=2, pady=2, sticky=W, row=2, column=0)

    kolumnaopisweza = Entry(matka, width=30)
    kolumnaopisweza.grid(padx=2, pady=2, sticky=E, row=2, column=1)

    Label(matka, text='Warstwa zewnętrzna:').grid(padx=2, pady=2, sticky=W, row=3, column=0)

    kolumnawarstwazewnetrzna = Entry(matka, width=30)
    kolumnawarstwazewnetrzna.grid(padx=2, pady=2, sticky=E, row=3, column=1)

    Label(matka, text='Wzmocnienie:').grid(padx=2, pady=2, sticky=W, row=4, column=0)

    kolumnawzmocnienie = Entry(matka, width=30)
    kolumnawzmocnienie.grid(padx=2, pady=2, sticky=E, row=4, column=1)

    Label(matka, text='Warstwa wewnętrzna:').grid(padx=2, pady=2, sticky=W, row=5, column=0)

    kolumnawarstwawewnetrzna = Entry(matka, width=30)
    kolumnawarstwawewnetrzna.grid(padx=2, pady=2, sticky=E, row=5, column=1)

    Label(matka, text='Temp. maksymalna (°C):').grid(padx=2, pady=2, sticky=W, row=6, column=0)

    kolumnatempmaksymalna = Entry(matka, width=30)
    kolumnatempmaksymalna.grid(padx=2, pady=2, sticky=E, row=6, column=1)

    Label(matka, text='Temp. minimalna (°C):').grid(padx=2, pady=2, sticky=W, row=7, column=0)

    kolumnatempminimalna = Entry(matka, width=30)
    kolumnatempminimalna.grid(padx=2, pady=2, sticky=E, row=7, column=1)

    Label(matka, text='Średnica wewnętrzna (mm):').grid(padx=2, pady=2, sticky=W, row=8, column=0)

    kolumnasrednica = Entry(matka, width=30)
    kolumnasrednica.grid(padx=2, pady=2, sticky=E, row=8, column=1)

    Label(matka, text='Ciśnienie robocze (bar):').grid(padx=2, pady=2, sticky=W, row=9, column=0)

    kolumnacisnienierobocze = Entry(matka, width=30)
    kolumnacisnienierobocze.grid(padx=2, pady=2, sticky=E, row=9, column=1)

    Label(matka, text='Ciśnienie rozerwania (bar):').grid(padx=2, pady=2, sticky=W, row=10, column=0)

    kolumnacisnienierozerwania = Entry(matka, width=30)
    kolumnacisnienierozerwania.grid(padx=2, pady=2, sticky=E, row=10, column=1)

    Label(matka, text='Promień zagięcia (statyczny) (mm):').grid(padx=2, pady=2, sticky=W, row=11, column=0)

    kolumnapromienstatyczny = Entry(matka, width=30)
    kolumnapromienstatyczny.grid(padx=2, pady=2, sticky=E, row=11, column=1)

    Label(matka, text='Promień zagięcia (dynamiczny) (mm):').grid(padx=2, pady=2, sticky=W, row=12, column=0)

    kolumnapromiendynamiczny = Entry(matka, width=30)
    kolumnapromiendynamiczny.grid(padx=2, pady=2, sticky=E, row=12, column=1)

    Label(matka, text='Flaga tolerancji:').grid(padx=2, pady=2, sticky=W, row=13, column=0)
    Button(matka, text='?', command=pomocflaga).grid(sticky=E, row=13, column=0)

    kolumnaflagatolerancji = ttk.Combobox(matka, state='readonly', width=27)
    kolumnaflagatolerancji['values'] = ['P1', 'P2', 'P3', 'S1', 'S2', 'S3', 'S4', 'S5', 'K', 'C']
    kolumnaflagatolerancji.grid(padx=2, pady=2, sticky=E, row=13, column=1)
    kolumnaflagatolerancji.current(0)

    Button(matka, text='P', command=lambda:otworztoler('P'), width=3).grid(row=14, padx=20, pady=6, sticky=W)
    Button(matka, text='S', command=lambda:otworztoler('S'), width=3).grid(row=14, padx=20, pady=6)
    Button(matka, text='K', command=lambda:otworztoler('K'), width=3).grid(row=14, padx=20, pady=6, sticky=E)
    Button(matka, text='C', command=lambda:otworztoler('C'), width=3).grid(row=14, padx=20, pady=6, column=1, sticky=W)

    Button(matka, text='Dodaj wąż!', height=2, command=dodajwazdolisty).grid(padx=2, pady=10, sticky=E, row=15, column=0, columnspan=1)

    matka.mainloop()

def pomocflaga():
    messagebox.showinfo('Pomoc nt. flag tolerancji', 'Flagi tolerancji to oznaczenie, dzięki któremu na świadectwie pojawia się informacja o tolerancji wykonania przewodu.\n\nKategorie:\nP - dla przewodów \"podstawowych\"\nS - dla przewodów stalowych\nK - dla przewodów kompozytowych\nC - dla przewodów Corroflon\n\nCyfry przy znacznikach P i S oznaczają kategorię ze względu na średnicę, tak jak podają instrukcje tolerancji dla tych rodzajów przewodów.\n\nW razie wątpliwości kontakt z Mariuszem Węcławiakiem z KJ.')

def checkifitsce():
    if typoznaczenia.get() == 'CE I (stalowy)' or typoznaczenia.get() == 'CE II (stalowy)' or typoznaczenia.get() == 'CE I (niestalowy)':
        return 1
    else:
        return 0

def sprawdzczywybranowaz():
    if tekstindeksuweza.get() != '':
        return 1
    else:
        messagebox.showwarning("Nie wybrano indeksu węża!", "Nie wybrano indeksu węża!\n\nJeżeli szukanego węża nie ma liście, wybierz pustą pozycję z listy (tzw. zapchajdziurę)!")
        return 0

def onselectoznakowanie(event):
    cowybrano = typoznaczenia.get()
    if cowybrano == 'NIE - oznaczenie standardowe':
        poleoznaczenia.delete(0.0, END)
    elif cowybrano == 'Pharmaline N/G':
        poleoznaczenia.delete(0.0, END)
        if "PHGP-X" in tekstindeksuweza.get():
            a = "X"
        else:
            a = "N"
        
        if "PHGP" in tekstindeksuweza.get():
            dn = int("".join(list(filter(str.isdigit, tekstindeksuweza.get()))))
        else:
            dn = "XX"

        poleoznaczenia.insert(1.0, 'PH'+a+' DN'+str(dn)+' PN'+cisnienieRob.get()+'\nTUBES INT. '+numerZP.get())
    elif cowybrano == 'CORROFLON':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'CFSS DNXX PN'+cisnienieRob.get()+'\nTUBES INT. '+numerZP.get())
    elif cowybrano == 'CE I (stalowy)':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'TI-'+str(now.year)+'-EN14585-1-DNXX-PS'+cisnienieRob.get()+'\nTS TEMP°C / TEMP°C-'+numerZP.get())
    elif cowybrano == 'CE II (stalowy)':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'TI-'+str(now.year)+'-EN14585-1-DNXX-PS'+cisnienieRob.get()+'\nTS TEMP°C / TEMP°C-'+numerZP.get())
    elif cowybrano == 'CE I (niestalowy)':
        if sprawdzczywybranowaz() == 1:
            wyplujdane()
            poleoznaczenia.delete(0.0, END)
            poleoznaczenia.insert(1.0, 'TUBES INTERNATIONAL\nDNXX '+listaparametrow[0]+ strftime(' %m/%y') + '\nPS/PT (BAR) ' + cisnienieRob.get() +'/'+cisnienieTest.get()+(' '+numerZP.get()))
        else:
            pass
    elif cowybrano == 'Spir Star':
        if sprawdzczywybranowaz() == 1:
            wyplujdane()
            poleoznaczenia.delete(0.0, END)
            numerzmyslnikiem = numerZP.get()
            numerbezmyslnika = numerzmyslnikiem.replace("-", "")
            modelweza = tekstindeksuweza.get()
            a = modelweza.replace("SS-NW-", "")
            b = a.replace("-", "/", 1)
            poleoznaczenia.insert(1.0, 'TUBES '+strftime('%m %y'+'\n')+numerbezmyslnika+' BNRXXXXX\nTYPE '+b+' WP '+str(int(listaparametrow[9]))+' BAR')
        else:
            pass
    elif cowybrano == 'COLGATE':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, strftime('%m.%Y\n')+tu.get()+'\n'+numerZP.get()+'\nwww.tubes-international.com')
    elif cowybrano == 'PESA (OEM)':
        poleoznaczenia.delete(0.0, END)
        poleoznaczenia.insert(1.0, 'Tutaj wpisać oznaczenie PESA')

def szukaj(name1, name2, op):
    wprowadzony = szukajweza.get()
    matching = [s for s in listawezy if wprowadzony.lower() in s.lower()]
    rezultaty.delete(0, END)
    rezultaty.insert(END, *matching)
    rezultaty.selection_anchor(0)
    wprowadzwybor()

def wprowadzwybor():
    tekstindeksuweza.set(rezultaty.get(ANCHOR))

def kliknijwybor(event):
    tekstindeksuweza.set(rezultaty.get(ANCHOR))

def wyplujdane():
    wybranywaz = tekstindeksuweza.get()
    del listaparametrow[:]
    for wiersz in range(ws.nrows):
        if wybranywaz == ws.cell_value(wiersz,0):
            wierszweza = wiersz
            for cell in ws.row(wierszweza):
                cell_value = cell.value
                if cell.ctype in (2,3) and int(cell_value) == cell_value:
                    cell_value = str(int(cell_value))
                elif cell.ctype in (2,3) and str(cell_value) == cell_value:
                    cell_value = str(cell_value)
                listaparametrow.append(cell_value)

def wyczyscpola():
    potwierdzeniewyczyszczenia = messagebox.askyesno('Potwierdź', 'Czy chcesz wyczyścić wszystkie pola formularza?')
    if potwierdzeniewyczyszczenia:
        podcisn.activate()
        badanieTak.select()
        naccheck()
        obiektywarunkowklienta= [medium, cisnienieRob, tempWew, tempZew, ilosc, dlugosc, podcisn, cisnienieTest, czasTestu, szukajweza, koncowka1entry, koncowka2entry, entryKlient1, entryKlient2, entryKlient3, numerQC, numerZP]
        for i in obiektywarunkowklienta:
            i.delete(0, END)
            try:
                i.validate()
            except AttributeError:
                pass
        czypodcisn.set(0)
        podcisn.disable()
        tu.delete(0, END)
        mediumBadania.current([0])
        numerZam.delete(0, END)
        tekstindeksuweza.set('')
        czyzuzycie.set(0)
        czyodtlu.set(0)
        czyciaglosc.set(0)
        poleuwagi.delete('0.0', END)
        poleoznaczenia.delete('0.0', END)
        typoznaczenia.current(0)
        szukajweza.delete(0, END)
        rezultaty.delete(0, END)
        tekstindeksuweza.set('')
        koncowki.current(0)
        tuleje.current(0)
        tulejwybor.current(0)
    else:
        pass

def generatortolerancji():
    flaga = listaparametrow[13]
    dlu = int(dlugosc.get())
    global toler
    if flaga == 'P1':
        if dlu <= 630:
            toler = '-3 mm / +7 mm'
        elif 630 < dlu <= 1250:
            toler = '-4 mm / +12 mm'
        elif 1250 < dlu <= 2500:
            toler = '-6 mm / +20 mm'
        elif 2500 < dlu <= 8000:
            toler = '-0,5% / +1,5%'
        elif 8000 < dlu:
            toler = '-1% / +3%'
    elif flaga == 'P2':
        if dlu <= 630:
            toler = '-4 mm / +12 mm'
        elif 630 < dlu <= 1250:
            toler = '-6 mm / +20 mm'
        elif 1250 < dlu <= 2500:
            toler = '-6 mm / +25 mm'
        elif 2500 < dlu <= 8000:
            toler = '-0,5% / +1,5%'
        elif 8000 < dlu:
            toler = '-1% / +3%'
    elif flaga == 'P3':
        if dlu <= 2500:
            toler = '-6 mm / +25 mm'
        elif 2500 < dlu <= 8000:
            toler = '-0,5% / +1,5%'
        elif 8000 < dlu:
            toler = '-1% / +3%'
    elif flaga == 'S1':
        if dlu <= 1000:
            toler = '+10 mm'
        elif dlu > 1000:
            toler = '+1%'
    elif flaga == 'S2':
        if dlu <= 1000:
            toler = '+15 mm'
        elif dlu > 1000:
            toler = '+1,5%'
    elif flaga == 'S3':
        if dlu <= 1000:
            toler = '+20 mm'
        elif dlu > 1000:
            toler = '+2%'
    elif flaga == 'S4':
        if dlu <= 1000:
            toler = '+30 mm'
        elif dlu > 1000:
            toler = '+3%'
    elif flaga == 'S5':
        if dlu <= 1000:
            toler = '+40 mm'
        elif dlu > 1000:
            toler = '+4%'
    elif flaga == 'K':
        if dlu <= 2500:
            toler = '± 50 mm'
        elif dlu > 2500:
            toler = '± 2%'
    elif flaga == 'C':
        if dlu <= 1000:
            toler = '+5%'
        elif dlu > 1000:
            toler = '+10%'
    return toler


class Tooltip:
    '''
    It creates a tooltips for a given widget as the mouse goes on it.

    see:

    http://stackoverflow.com/questions/3221956/
           what-is-the-simplest-way-to-make-tooltips-
           in-tkinter/36221216#36221216

    http://www.daniweb.com/programming/software-development/
           code/484591/a-tooltips-class-for-tkinter

    - Originally written by vegaseat on 2014.09.09.

    - Modified to include a delay time by Victor Zaccardo on 2016.03.25.

    - Modified
        - to correct extreme right and extreme bottom behavior,
        - to stay inside the screen whenever the tooltips might go out on
          the top but still the screen is higher than the tooltips,
        - to use the more flexible mouse positioning,
        - to add customizable background color, padding, waittime and
          wraplength on creation
      by Alberto Vassena on 2016.11.05.

      Tested on Ubuntu 16.04/16.10, running Python 3.5.2

    TOD: themes styles support
    '''

    def __init__(self, widget,
                 *,
                 bg='#FFFFCA',
                 pad=(5, 3, 5, 3),
                 text='widget info',
                 waittime=1000,
                 wraplength=400):

        self.waittime = waittime  # in miliseconds, originally 500
        self.wraplength = wraplength  # in pixels, originally 180
        self.widget = widget
        self.text = text
        self.widget.bind("<Enter>", self.onEnter)
        self.widget.bind("<Leave>", self.onLeave)
        self.widget.bind("<ButtonPress>", self.onLeave)
        self.bg = bg
        self.pad = pad
        self.id = None
        self.tw = None
    
    def onEnter(self, event=None):
        if wlaczniktooltip.get() == False:
            pass
        else:
            self.schedule()

    def onLeave(self, event=None):
        self.unschedule()
        self.hide()

    def schedule(self):
        self.unschedule()
        self.id = self.widget.after(self.waittime, self.show)

    def unschedule(self):
        id_ = self.id
        self.id = None
        if id_:
            self.widget.after_cancel(id_)

    def show(self):
        def tip_pos_calculator(widget, label,
                               *,
                               tip_delta=(10, 5), pad=(5, 3, 5, 3)):

            w = widget

            s_width, s_height = w.winfo_screenwidth(), w.winfo_screenheight()

            width, height = (pad[0] + label.winfo_reqwidth() + pad[2],
                             pad[1] + label.winfo_reqheight() + pad[3])

            mouse_x, mouse_y = w.winfo_pointerxy()

            x1, y1 = mouse_x + tip_delta[0], mouse_y + tip_delta[1]
            x2, y2 = x1 + width, y1 + height

            x_delta = x2 - s_width
            if x_delta < 0:
                x_delta = 0
            y_delta = y2 - s_height
            if y_delta < 0:
                y_delta = 0

            offscreen = (x_delta, y_delta) != (0, 0)

            if offscreen:

                if x_delta:
                    x1 = mouse_x - tip_delta[0] - width

                if y_delta:
                    y1 = mouse_y - tip_delta[1] - height

            offscreen_again = y1 < 0  # out on the top

            if offscreen_again:
                # No further checks will be done.

                # TIP:
                # A further mod might automagically augment the
                # wraplength when the tooltips is too high to be
                # kept inside the screen.
                y1 = 0

            return x1, y1

        bg = self.bg
        pad = self.pad
        widget = self.widget

        # creates a toplevel window
        self.tw = Toplevel(widget)

        # Leaves only the label and removes the app window
        self.tw.wm_overrideredirect(True)

        win = Frame(self.tw,
                       background=bg,
                       borderwidth=0)
        label = Label(win,
                          text=self.text,
                          justify=LEFT,
                          background=bg,
                          relief=SOLID,
                          borderwidth=0,
                          wraplength=self.wraplength)

        label.grid(padx=(pad[0], pad[2]),
                   pady=(pad[1], pad[3]),
                   sticky=NSEW)
        win.grid()

        x, y = tip_pos_calculator(widget, label)

        self.tw.wm_geometry("+%d+%d" % (x, y))

    def hide(self):
        tw = self.tw
        if tw:
            tw.destroy()
        self.tw = None


def sprawdzczyjestdlugosc():
    if dlugosc.get().isdigit() == True and ilosc.get().isdigit() == True:
        return 1
    else:
        messagebox.showwarning("Nie wpisano prawidłowej długości/ilości!", "Upewnij się, że w polach ilości oraz długości przewodu/węża znajdują się prawidłowe wartości!")
        return 0

def sprawdzenieprzedzapisem():
    if checkifitsce() == 1:
        if tempWew.get() == 'otoczenie' or tempWew.get() == 'otoczenia' or tempZew.get() == 'otoczenie' or tempZew.get() == 'otoczenia':
            messagebox.showwarning('Nieprawidłowa temperatura!', 'Przy oznakowaniu CE należy podać liczbowy przedział temperatury!')
            return
        else:
            pass
    else:
        pass
    
    if sprawdzczywybranowaz() == 1 and sprawdzczyjestdlugosc() == 1:
        if wersjajezykowa.get() == 'PL':
            generujpopolsku()
        elif wersjajezykowa.get() == 'ENG':
            generujpoangielsku()
        else:
            messagebox.showwarning('Coś poszło nie tak!', 'Coś poszło nie tak z generowaniem dokumentu. Spróbuj ponownie lub zgłoś błąd!')
    else:
        pass

def generujpopolsku():
    wyplujdane()
    generatortolerancji()
    global checkIle, checkIle2
    nQC = numerQC.get()
    nKRP = numerZP.get()
    nTU = tu.get()
    ile = int(ilosc.get())
    rodzdlu = waz.get()
    dlu = dlugosc.get()
    med = medium.get()
    cisR = cisnienieRob.get()
    pdcsn = podcisn.get().replace('.',',')
    tw = tempWew.get()
    tz = tempZew.get()
    medB = mediumBadania.get()
    cisT = cisnienieTest.get()
    cz = czasTestu.get()
    zam = numerZam.get()
    przygotowal = entryprzygotowal.get()
    kiedy_p = entrykiedyprzygotowal.get()
    zatwierdzil = entryzatwierdził.get()
    kiedy_z = entrykiedyzatwierdzil.get()
    kli1 = entryKlient1.get()
    kli2 = entryKlient2.get()
    kli3 = entryKlient3.get()
    konc1 = koncowka1entry.get()
    konc2 = koncowka2entry.get()
    uwagi = poleuwagi.get(0.0, "end-1c")
    kol1 = listaparametrow[0]
    kol2 = listaparametrow[1]
    kol3 = listaparametrow[2]
    kol4 = listaparametrow[3]
    kol5 = listaparametrow[4]
    kol6 = listaparametrow[5]
    kol7 = listaparametrow[6]
    kol8 = listaparametrow[7]
    kol9 = listaparametrow[8]
    kol10 = listaparametrow[9]
    kol11 = listaparametrow[10]
    kol12 = listaparametrow[11]
    kol13 = listaparametrow[12]
    wyboroznak = typoznaczenia.get()

    if ile == 1:
        checkIle = 'przewód'
    elif ile > 4:
        checkIle = 'przewodów'
    elif ile == 2 or 3 or 4:
        checkIle = 'przewody'
    else:
        print('errorrrr')

    if ile == 1:
        checkIle2 = 'Przewód'
    elif ile > 1:
        checkIle2 = 'Przewody'
    else:
        print('errrororor2')

    if ktoryoddzial.get() == '50':
        numeroddzialu = ''
    else:
        numeroddzialu = ktoryoddzial.get()+'/'


    if typoznaczenia.get() == 'CE I (stalowy)' or typoznaczenia.get() == 'CE II (stalowy)' or typoznaczenia.get() == 'CE I (niestalowy)':
        document = Document('bin\wzor_ce.docx')
        firstparagraph = document.paragraphs[0]
        firstparagraph.add_run('ŚWIADECTWO JAKOŚCI\tQC/'+numeroddzialu+nQC+strftime('/%y'))
        firstparagraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        firstparagraph.style = 'Heading 1'
    else:
        document = Document('bin\wzor2.docx')
        firstparagraph = document.paragraphs[0]
        firstparagraph.add_run('ŚWIADECTWO JAKOŚCI\tQC/'+numeroddzialu + nQC+strftime('/%y'))
        firstparagraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        firstparagraph.style = 'Heading 1'

    klient = document.add_paragraph('Wystawiono dla:\t'+kli1)
    klient.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph('\t'+kli2).paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph('\t'+kli3).paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph()
    if zam == '':
        pass
    else:
        zamowienie = document.add_paragraph('Zamówienie:\t' + zam)
        zamowienie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()

    wykonanie = document.add_paragraph('Wykonanie:\t'+checkIle2+' wykonano wg zlecenia produkcyjnego ' + nKRP)
    wykonanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph()

    if wyboroznak == 'NIE - oznaczenie standardowe':
        oznakowanie = document.add_paragraph('Oznakowanie:\t'+checkIle2+' oznakowano ' + nKRP)
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()
    elif wyboroznak == 'Pharmaline N/G':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t'+poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t'+poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER)  # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'CORROFLON':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t'+poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t'+poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'CE I (stalowy)':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'CE II (stalowy)':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t1433\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(10.25), WD_TAB_ALIGNMENT.LEFT) # @UndefinedVariable
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    elif wyboroznak == 'CE I (niestalowy)':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'Spir Star':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'COLGATE':
        oznakowanie = document.add_paragraph('Oznakowanie:\t' + checkIle2 + ' oznakowano\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie4 = document.add_paragraph('\t' + poleoznaczenia.get(4.0, '4.end'))
        oznakowanie4.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'PESA (OEM)':
        oznakowanie = document.add_paragraph('Oznakowanie:\t'+checkIle2+' oznakowano ' + nKRP)
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie2 = document.add_paragraph(poleoznaczenia.get(1.0, "1.end"))
        oznakowanie2.paragraph_format.left_indent = (Cm(4))
        document.add_paragraph()
        
    opisprzewodu = document.add_paragraph(
        'Wyrób:\t'+checkIle2+ ' '+nTU + ' wykonano z węża typu '+kol2+' Ø '+str(kol9).replace('.',',')+' mm ('+kol1+'), zakończono')
    opisprzewodu.paragraph_format.left_indent = (Cm(4))
    opisprzewodu.paragraph_format.first_line_indent = (Cm(-4))
    opisprzewodu.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    opisprzewodu.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

    if ilekoncowek.get() == 1:
        opisprzewodu.add_run(' końcówkami:')
        koncowka1 = document.add_paragraph('Końcówka 1: '+konc1)
        koncowka1.paragraph_format.left_indent = (Cm(6.25))
        koncowka1.paragraph_format.first_line_indent = (Cm(-2.25))
        koncowka1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

        koncowka2 = document.add_paragraph('Końcówka 2: '+konc2)
        koncowka2.paragraph_format.left_indent = (Cm(6.25))
        koncowka2.paragraph_format.first_line_indent = (Cm(-2.25))
        koncowka2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable
    elif ilekoncowek.get() == 2:
        opisprzewodu.add_run(' obustronnie końcówkami '+konc1+'.')
    else:
        pass

    if koncowki.get() == tuleje.get():
        materialkoncowek = document.add_paragraph('Końcówki i '+tulejwybor.get()+' zostały wykonane '+listamaterialow.get(koncowki.get())+'.')
    else:
        materialkoncowek = document.add_paragraph('Końcówki zostały wykonane '+listamaterialow.get(koncowki.get())+'. '+tulejwybor.get().title()+' zostały wykonane '+listamaterialow.get(tuleje.get())+'.')
    materialkoncowek.paragraph_format.left_indent = (Cm(4))
    materialkoncowek.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

    if uwagi == '':
        pass
    else:
        xyz = uwagi.split("\n")
        for i in xyz:
            wpisaneuwagi = document.add_paragraph(i)
            wpisaneuwagi.paragraph_format.left_indent = (Cm(4))
            wpisaneuwagi.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

    document.add_paragraph()

    opisweza = document.add_paragraph(kol3)
    opisweza.paragraph_format.left_indent = (Cm(4))
    opisweza.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable
    document.add_paragraph()

    if kol6 == '':
        pass
    else:
        wwew = document.add_paragraph('Warstwa wewn.:\t'+kol6)
        wwew.paragraph_format.left_indent = (Cm(4))

    if kol5 == '':
        pass
    else:
        wzmoc = document.add_paragraph('Wzmocnienie:\t'+kol5)
        wzmoc.paragraph_format.left_indent = (Cm(4))

    if kol4 == '':
        pass
    else:
        wzew = document.add_paragraph('Warstwa zewn.:\t'+kol4)
        wzew.paragraph_format.left_indent = (Cm(4))

    document.add_paragraph()

    if rodzdlu == 'Długość przewodu (mm)':
        ilo = document.add_paragraph('Ilość:\t' + str(ile) + ' ' + checkIle + ' o długości ' + dlu + ' mm. Tolerancja długości '+toler+'.')
        ilo.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()
    else:
        ilo = document.add_paragraph('Ilość:\t' + str(ile) + ' ' + checkIle + ' o długości węża ' + dlu + ' mm. Tolerancja długości '+toler+'.')
        ilo.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()

    if kol13 != '':
        document.add_paragraph('Katalogowe parametry pracy węża:')

        parametryweza0 = document.add_paragraph('\twarunki statyczne / dynamiczne')
        parametryweza0.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza0.paragraph_format.left_indent = (Cm(4))

        parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
        parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza1.paragraph_format.left_indent = (Cm(4))

        parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t' + str(kol11) + ' bar')
        parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza2.paragraph_format.left_indent = (Cm(4))

        parametryweza3 = document.add_paragraph('- promień zagięcia:\t' + str(kol12) + ' mm / '+str(kol13)+' mm')
        parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza3.paragraph_format.left_indent = (Cm(4))

        parametryweza4 = document.add_paragraph('- temperatura pracy:\tod ' + str(int(kol8)) + '°C do +' + str(int(kol7)) + '°C')
        parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza4.paragraph_format.left_indent = (Cm(4))

    else:
        if kol1 == ' ': #pusty indeks
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t ')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- promień zagięcia:\t')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph('- temperatura pracy:\t')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        elif kol11 == '': #brak w katalogu ciśnienia rozerwania
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- promień zagięcia:\t' + str(kol12) + ' mm')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperatura pracy:\tod ' + str(kol8) + '°C do +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        elif kol12 == '': #brak w katalogu promienia zagięcia
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t' + str(kol11) + ' bar')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperatura pracy:\tod ' + str(kol8) + '°C do +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        else: #wszystko jest
            document.add_paragraph('Katalogowe parametry pracy węża:')
            parametryweza1 = document.add_paragraph('- maksymalne ciśnienie robocze:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- ciśnienie rozerwania:\t' + str(kol11) + ' bar')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- promień zagięcia:\t' + str(kol12) + ' mm')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperatura pracy:\tod ' + str(kol8) + '°C do +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))


    document.add_paragraph()
    document.add_paragraph('Parametry pracy określone przez zamawiającego:')

    parametryklienta2 = document.add_paragraph('- medium:\t'+med)
    parametryklienta2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryklienta2.paragraph_format.left_indent = (Cm(4))

    if checkifitsce() == 1:
        parametryklienta1 = document.add_paragraph('- ciśnienie robocze PS:\t'+cisR+' bar')
    else:
        parametryklienta1 = document.add_paragraph('- ciśnienie robocze:\t'+cisR+' bar')
    parametryklienta1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryklienta1.paragraph_format.left_indent = (Cm(4))

    if podcisn.get() != '':
        parametryklienta5 = document.add_paragraph('- podciśnienie:\t'+pdcsn+' bar')
        parametryklienta5.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryklienta5.paragraph_format.left_indent = (Cm(4))
    else:
        pass

    if tw == 'otoczenia' or tw == 'otoczenie':
        if checkifitsce() == 1:
            parametryklienta3 = document.add_paragraph('- temp. wewnętrzna TS:\totoczenia')
        else:
            parametryklienta3 = document.add_paragraph('- temp. wewnętrzna:\totoczenia')
    elif tw == '':
        parametryklienta3 = document.add_paragraph('- temp. wewnętrzna:\tbrak danych')
    else:
        if checkifitsce() == 1:
            parametryklienta3 = document.add_paragraph('- temp. wewnętrzna TS:\t' + tw + '°C')
        else:
            parametryklienta3 = document.add_paragraph('- temp. wewnętrzna:\t' + tw + '°C')
    parametryklienta3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryklienta3.paragraph_format.left_indent = (Cm(4))

    if tz == 'otoczenia' or tz == 'otoczenie':
        if checkifitsce() == 1:
            parametryklienta4 = document.add_paragraph('- temp. zewnętrzna TS:\totoczenia')
        else:
            parametryklienta4 = document.add_paragraph('- temp. zewnętrzna:\totoczenia')
    elif tz =='':
        parametryklienta4 = document.add_paragraph('- temp. zewnętrzna:\tbrak danych')
    else:
        if checkifitsce() == 1:
            parametryklienta4 = document.add_paragraph('- temp. zewnętrzna TS:\t' + tz + '°C')
        else:
            parametryklienta4 = document.add_paragraph('- temp. zewnętrzna:\t' + tz + '°C')
    parametryklienta4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryklienta4.paragraph_format.left_indent = (Cm(4))

    if checkifitsce() == 1:
        cisnienieproby = document.add_paragraph('- ciśnienie próby PT:\t'+ cisT + ' bar')
        cisnienieproby.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        cisnienieproby.paragraph_format.left_indent = (Cm(4))
    else:
        pass
    document.add_paragraph()

    czytest2 = czytest.get()
    if czytest2 == 1:
        if cz == '1':
            ktoraminuta = 'minutę'
        elif cz == '2' or cz == '3' or cz == '4' or cz == '0,5':
            ktoraminuta = 'minuty'
        else:
            ktoraminuta = 'minut'
        teksttestu1 = document.add_paragraph()
        teksttestu1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        teksttestu1.add_run(checkIle2+' testowano ' + medB + ', pod ciśnieniem ' + cisT + ' bar przez ' + cz + ' '+ktoraminuta+'.').font.underline = TRUE

        teksttestu2 = document.add_paragraph()
        teksttestu2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        teksttestu2.add_run('Nieszczelności nie stwierdzono.').font.underline = TRUE
        document.add_paragraph()
    else:
        pass

    if czyodtlu.get() == 1:
        tekstodtluszczania = document.add_paragraph()
        tekstodtluszczania.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        tekstodtluszczania.add_run(checkIle2+' poddano procesowi wstępnego odtłuszczania do pracy z tlenem.').font.underline = TRUE
        tekstodtluszczania2 = document.add_paragraph()
        tekstodtluszczania2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        tekstodtluszczania2.add_run('Odtłuszczenie finalne i dopuszczenie do eksploatacji leży po stronie zamawiającego.').font.underline = TRUE
        document.add_paragraph()
    else:
        pass

    if czyzuzycie.get() == 1:
        if ile == 1:
            tekstzuzycia = document.add_paragraph()
            tekstzuzycia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
            tekstzuzycia.add_run('Ze względu na różnorodność mediów oraz ich agresywność, żywotność przewodu może być ograniczona.').font.underline = TRUE
            document.add_paragraph()
        else:
            tekstzuzycia = document.add_paragraph()
            tekstzuzycia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
            tekstzuzycia.add_run('Ze względu na różnorodność mediów oraz ich agresywność, żywotność przewodów może być ograniczona.').font.underline = TRUE
            document.add_paragraph()
    else:
        pass

    if czyciaglosc.get() == 1:
        if ile == 1:
            tekstciaglosc = document.add_paragraph()
            tekstciaglosc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
            tekstciaglosc.add_run(
                'Potwierdzono badaniem zachowanie ciągłości elektrycznej pomiędzy końcówkami przewodu.').font.underline = TRUE
            document.add_paragraph()
        else:
            tekstciaglosc = document.add_paragraph()
            tekstciaglosc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
            tekstciaglosc.add_run(
                'Potwierdzono badaniem zachowanie ciągłości elektrycznej pomiędzy końcówkami przewodów.').font.underline = TRUE
            document.add_paragraph()
    else:
        pass

    if ile == 1:
        document.add_paragraph('Przewód może być użytkowany w parametrach określonych przez zamawiającego.')
    else:
        document.add_paragraph('Przewody mogą być użytkowane w parametrach określonych przez zamawiającego.')
    document.add_paragraph()

    if ile == 1:
        uzytkowanie1 = document.add_paragraph('Użytkowanie przewodu wg stanu technicznego - biorąc pod uwagę mechaniczne zużycie, ewentualne przetarcia podczas użytkowania, stopień ryzyka dla obsługi i wszystkie inne czynniki mogące mieć wpływ na zastosowanie i żywotność przewodów.')
        uzytkowanie1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable
    else:
        uzytkowanie1 = document.add_paragraph('Użytkowanie przewodów wg stanu technicznego - biorąc pod uwagę mechaniczne zużycie, ewentualne przetarcia podczas użytkowania, stopień ryzyka dla obsługi i wszystkie inne czynniki mogące mieć wpływ na zastosowanie i żywotność przewodów.')
        uzytkowanie1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable
    document.add_paragraph()

    if ktoryoddzial.get() == '50':
        przygzatw = document.add_paragraph('Przygotował:\t' + przygotowal + '\tZatwierdził:\t' + zatwierdzil)
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(13.75))

        datydaty = document.add_paragraph('Data:\t' + kiedy_p + '\tData:\t' + kiedy_z)
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(13.75))

        podpispodpis = document.add_paragraph('Podpis:\tPodpis:')
        podpispodpis.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
        podpispodpis.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
    else:
        przygzatw = document.add_paragraph('Przygotował:\t' + przygotowal)
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))

        datydaty = document.add_paragraph('Data:\t' + kiedy_p)
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))

        podpispodpis = document.add_paragraph('Podpis:')

    def filesave():
        if ktoryoddzial.get() == '50':
            sss = nQC
        else:
            sss = 'QC-'+ktoryoddzial.get()+'-'+nQC+'-'+strftime('%y')
        filename = filedialog.asksaveasfilename(defaultextension='.docx', initialfile=sss)
        if not filename: return
        document.save(filename)
        pytanieootwarcie = messagebox.askyesno("Co teraz?", "Gratuluję sukcesu!\nCzy chcesz teraz otworzyć zapisany dokument?", default="yes")
        if pytanieootwarcie:
            startfile(filename)
        else:
            pass
    filesave()

def generujpoangielsku():
    wyplujdane()
    generatortolerancji()
    global checkIle, checkIle2
    nQC = numerQC.get()
    nKRP = numerZP.get()
    nTU = tu.get()
    ile = int(ilosc.get())
    rodzdlu = waz.get()
    dlu = dlugosc.get()
    med = medium.get()
    cisR = cisnienieRob.get()
    pdcsn = podcisn.get().replace('.',',')
    tw = tempWew.get()
    tz = tempZew.get()
    medB = mediumBadania.get()
    cisT = cisnienieTest.get()
    cz = czasTestu.get()
    zam = numerZam.get()
    medben = ''
    hashave = 'water tested'
    przygotowal = entryprzygotowal.get()
    kiedy_p = entrykiedyprzygotowal.get()
    zatwierdzil = entryzatwierdził.get()
    kiedy_z = entrykiedyzatwierdzil.get()
    kli1 = entryKlient1.get()
    kli2 = entryKlient2.get()
    kli3 = entryKlient3.get()
    konc1 = koncowka1entry.get()
    konc2 = koncowka2entry.get()
    uwagi = poleuwagi.get(0.0, "end-1c")
    kol1 = listaparametrow[0]
    kol2 = listaparametrow[1]
    kol3 = listaparametrow[2]
    kol4 = listaparametrow[3]
    kol5 = listaparametrow[4]
    kol6 = listaparametrow[5]
    kol7 = listaparametrow[6]
    kol8 = listaparametrow[7]
    kol9 = listaparametrow[8]
    kol10 = listaparametrow[9]
    kol11 = listaparametrow[10]
    kol12 = listaparametrow[11]
    kol13 = listaparametrow[12]
    wyboroznak = typoznaczenia.get()

    if ile == 1:
        checkIle = 'hose assembly'
    elif ile > 1:
        checkIle = 'hose assemblies'
    else:
        print('errorrrr')

    if ile == 1:
        checkIle2 = 'Hose assembly'
    elif ile > 1:
        checkIle2 = 'Hose assemblies'
    else:
        print('errrororor2')

    if ktoryoddzial.get() == '50':
        numeroddzialu = ''
    else:
        numeroddzialu = '/'+ktoryoddzial.get()+'/'

    if typoznaczenia.get() == 'CE I (stalowy)' or typoznaczenia.get() == 'CE II (stalowy)' or typoznaczenia.get() == 'CE I (niestalowy)':
        document = Document('bin\wzor_ce_en.docx')
        firstparagraph = document.paragraphs[0]
        firstparagraph.add_run('QUALITY CERTIFICATE\tQC/'+numeroddzialu+nQC+strftime('/%y'))
        firstparagraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        firstparagraph.style = 'Heading 1'
    else:
        document = Document('bin\wzor_en.docx')
        firstparagraph = document.paragraphs[0]
        firstparagraph.add_run('QUALITY CERTIFICATE\tQC/' + numeroddzialu+nQC+strftime('/%y'))
        firstparagraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        firstparagraph.style = 'Heading 1'

    klient = document.add_paragraph('Issued for:\t'+kli1)
    klient.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph('\t'+kli2).paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph('\t'+kli3).paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph()
    if zam == '':
        pass
    else:
        zamowienie = document.add_paragraph('Order:\t' + zam)
        zamowienie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()

    wykonanie = document.add_paragraph('Performance:\t'+checkIle2+' made according to job card ' + nKRP)
    wykonanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    document.add_paragraph()

    if wyboroznak == 'NIE - oznaczenie standardowe':
        oznakowanie = document.add_paragraph('Marking:\t'+checkIle2+' marked ' + nKRP)
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()
    elif wyboroznak == 'Pharmaline N/G':
        oznakowanie = document.add_paragraph('Marking:\t' + checkIle2 + ' marked\t'+poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t'+poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'CORROFLON':
        oznakowanie = document.add_paragraph('Marking:\t' + checkIle2 + ' marked\t'+poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t'+poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(12), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'CE I (stalowy)':
        oznakowanie = document.add_paragraph('Marking:\t' + checkIle2 + ' marked\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'CE II (stalowy)':
        oznakowanie = document.add_paragraph('Marking:\t' + checkIle2 + ' marked\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t1433\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(10.25), WD_TAB_ALIGNMENT.LEFT) # @UndefinedVariable
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    elif wyboroznak == 'CE I (niestalowy)':
        oznakowanie = document.add_paragraph('Marking:\t' + checkIle2 + ' marked\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'Spir Star':
        oznakowanie = document.add_paragraph('Marking:\t' + checkIle2 + ' marked\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'COLGATE':
        oznakowanie = document.add_paragraph('Marking:\t' + checkIle2 + ' marked\t' + poleoznaczenia.get(1.0, "1.end"))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie2 = document.add_paragraph('\t' + poleoznaczenia.get(2.0, '2.end'))
        oznakowanie2.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie3 = document.add_paragraph('\t' + poleoznaczenia.get(3.0, '3.end'))
        oznakowanie3.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        oznakowanie4 = document.add_paragraph('\t' + poleoznaczenia.get(4.0, '4.end'))
        oznakowanie4.paragraph_format.tab_stops.add_tab_stop(Cm(14), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        document.add_paragraph()
    elif wyboroznak == 'PESA (OEM)':
        oznakowanie = document.add_paragraph('Marking:\t'+checkIle2+' marked ' + nKRP)
        oznakowanie.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        oznakowanie2 = document.add_paragraph(poleoznaczenia.get(1.0, "1.end"))
        oznakowanie2.paragraph_format.left_indent = (Cm(4))
        document.add_paragraph()

    opisprzewodu = document.add_paragraph(
        'Item:\t'+checkIle2+' ' + nTU + ' made of '+kol2+' hose Ø '+str(kol9).replace('.',',')+' mm ('+kol1+'), ')
    opisprzewodu.paragraph_format.left_indent = (Cm(4))
    opisprzewodu.paragraph_format.first_line_indent = (Cm(-4))
    opisprzewodu.paragraph_format.tab_stops.add_tab_stop(Cm(4))
    opisprzewodu.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

    if ilekoncowek.get() == 1:
        opisprzewodu.add_run('ended with fittings:')
        koncowka1 = document.add_paragraph('Fitting 1: '+konc1)
        koncowka1.paragraph_format.left_indent = (Cm(6.25))
        koncowka1.paragraph_format.first_line_indent = (Cm(-2.25))
        koncowka1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

        koncowka2 = document.add_paragraph('Fitting 2: '+konc2)
        koncowka2.paragraph_format.left_indent = (Cm(6.25))
        koncowka2.paragraph_format.first_line_indent = (Cm(-2.25))
        koncowka2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable
    elif ilekoncowek.get() == 2:
        opisprzewodu.add_run('both sides ended with '+konc1+'.')
    else:
        pass
    
    materialypoang = {
        'tuleje': 'ferrules',
        'obejmy': 'clamps'
        }
    if koncowki.get() == tuleje.get():
        materialkoncowek = document.add_paragraph('Fittings and '+materialypoang.get(tulejwybor.get())+' made of '+listamaterialowang.get(koncowki.get())+'.')
    else:
        materialkoncowek = document.add_paragraph('Fittings made of '+listamaterialowang.get(koncowki.get())+'. '+materialypoang.get(tulejwybor.get()).title()+' made of '+listamaterialowang.get(tuleje.get())+'.')
    materialkoncowek.paragraph_format.left_indent = (Cm(4))
    materialkoncowek.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

    if uwagi == '':
        pass
    else:
        xyz = uwagi.split("\n")
        for i in xyz:
            wpisaneuwagi = document.add_paragraph(i)
            wpisaneuwagi.paragraph_format.left_indent = (Cm(4))
            wpisaneuwagi.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

    document.add_paragraph()

    opisweza = document.add_paragraph(kol3)
    opisweza.paragraph_format.left_indent = (Cm(4))
    opisweza.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable
    document.add_paragraph()

    if kol6 == '':
        pass
    else:
        wwew = document.add_paragraph('Internal layer:\t'+kol6)
        wwew.paragraph_format.left_indent = (Cm(4))

    if kol5 == '':
        pass
    else:
        wzmoc = document.add_paragraph('Reinforcement:\t'+kol5)
        wzmoc.paragraph_format.left_indent = (Cm(4))

    if kol4 == '':
        pass
    else:
        wzew = document.add_paragraph('External layer:\t'+kol4)
        wzew.paragraph_format.left_indent = (Cm(4))

    document.add_paragraph()

    if rodzdlu == 'Długość przewodu (mm)':
        ilo = document.add_paragraph('Quantity:\t' + str(ile) + ' ' + checkIle + ' - length ' + dlu + ' mm. Manufacturing tolerance '+toler+'.')
        ilo.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()
    else:
        ilo = document.add_paragraph('Quantity:\t' + str(ile) + ' ' + checkIle + ' - hose length ' + dlu + ' mm. Manufacturing tolerance '+toler+'.')
        ilo.paragraph_format.tab_stops.add_tab_stop(Cm(4))
        document.add_paragraph()

    if kol13 != '':
        document.add_paragraph('Catalogue working parameters:')

        parametryweza0 = document.add_paragraph('\tstatic / dynamic conditions')
        parametryweza0.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza0.paragraph_format.left_indent = (Cm(4))

        parametryweza1 = document.add_paragraph('- maximum working pressure:\t' + str(kol10) + ' bar')
        parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza1.paragraph_format.left_indent = (Cm(4))

        parametryweza2 = document.add_paragraph('- burst pressure:\t' + str(kol11) + ' bar')
        parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza2.paragraph_format.left_indent = (Cm(4))

        parametryweza3 = document.add_paragraph('- minimum bending radius:\t' + str(kol12) + ' mm / '+str(kol13)+' mm')
        parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza3.paragraph_format.left_indent = (Cm(4))

        parametryweza4 = document.add_paragraph('- temperature:\tfrom ' + str(int(kol8)) + '°C up to +' + str(int(kol7)) + '°C')
        parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryweza4.paragraph_format.left_indent = (Cm(4))

    else:
        if kol1 == ' ': #pusty indeks
            document.add_paragraph('Catalogue working parameters:')
            parametryweza1 = document.add_paragraph('- maximum working pressure:\t ')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- burst pressure:\t')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- minimum bending radius:\t')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph('- temperature:\t')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        elif kol11 == '': #brak w katalogu ciśnienia rozerwania
            document.add_paragraph('Catalogue working parameters:')
            parametryweza1 = document.add_paragraph('- maximum working pressure:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- minimum bending radius:\t' + str(kol12) + ' mm')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperature:\tfrom ' + str(kol8) + '°C up to +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        elif kol12 == '': #brak w katalogu promienia zagięcia
            document.add_paragraph('Catalogue working parameters:')
            parametryweza1 = document.add_paragraph('- maximum working pressure:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- burst pressure:\t' + str(kol11) + ' bar')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperature:\tfrom ' + str(kol8) + '°C up to +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))
        else: #wszystko jest
            document.add_paragraph('Catalogue working parameters:')
            parametryweza1 = document.add_paragraph('- maximum working pressure:\t' + str(kol10) + ' bar')
            parametryweza1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza1.paragraph_format.left_indent = (Cm(4))

            parametryweza2 = document.add_paragraph('- burst pressure:\t' + str(kol11) + ' bar')
            parametryweza2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza2.paragraph_format.left_indent = (Cm(4))

            parametryweza3 = document.add_paragraph('- minimum bending radius:\t' + str(kol12) + ' mm')
            parametryweza3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza3.paragraph_format.left_indent = (Cm(4))

            parametryweza4 = document.add_paragraph(
                '- temperature:\tfrom ' + str(kol8) + '°C up to +' + str(kol7) + '°C')
            parametryweza4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
            parametryweza4.paragraph_format.left_indent = (Cm(4))


    document.add_paragraph()
    document.add_paragraph('Working parameters stated by customer:')

    parametryklienta2 = document.add_paragraph('- medium:\t'+med)
    parametryklienta2.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryklienta2.paragraph_format.left_indent = (Cm(4))

    if checkifitsce() == 1:
        parametryKlienta1 = document.add_paragraph('- working pressure PS:\t'+cisR+' bar')
    else:
        parametryKlienta1 = document.add_paragraph('- working pressure:\t'+cisR+' bar')
    parametryKlienta1.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryKlienta1.paragraph_format.left_indent = (Cm(4))

    if podcisn.get() != '':
        parametryklienta5 = document.add_paragraph('- vacuum:\t'+pdcsn+' bar')
        parametryklienta5.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        parametryklienta5.paragraph_format.left_indent = (Cm(4))
    else:
        pass

    if tw == 'ambient' or tw == 'otoczenia' or tw == 'otoczenie':
        if checkifitsce() == 1:
            parametryklienta3 = document.add_paragraph('- internal temperature TS:\tambient')
        else:
            parametryklienta3 = document.add_paragraph('- internal temperature:\tambient')
    elif tw == '':
        parametryklienta3 = document.add_paragraph('- internal temperature:\tn/a')
    else:
        if checkifitsce() == 1:
            parametryklienta3 = document.add_paragraph('- internal temperature TS:\t' + tw + '°C')
        else:
            parametryklienta3 = document.add_paragraph('- internal temperature:\t' + tw + '°C')
    parametryklienta3.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryklienta3.paragraph_format.left_indent = (Cm(4))

    if tz == 'ambient' or tz == 'otoczenia' or tz == 'otoczenie':
        if checkifitsce() == 1:
            parametryklienta4 = document.add_paragraph('- external temperature:\tambient')
        else:
            parametryklienta4 = document.add_paragraph('- external temperature:\tambient')
    elif tz == '':
        parametryklienta4 = document.add_paragraph('- temp. zewnętrzna:\tn/a')
    else:
        if checkifitsce() == 1:
            parametryklienta4 = document.add_paragraph('- external temperature TS:\t' + tz + '°C')
        else:
            parametryklienta4 = document.add_paragraph('- external temperature:\t' + tz + '°C')
    parametryklienta4.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
    parametryklienta4.paragraph_format.left_indent = (Cm(4))

    if checkifitsce() == 1:
        cisnienieproby = document.add_paragraph('- test pressure PT:\t'+ cisT + ' bar')
        cisnienieproby.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.CENTER) # @UndefinedVariable
        cisnienieproby.paragraph_format.left_indent = (Cm(4))
    else:
        pass
    document.add_paragraph()

    czytest2 = czytest.get()
    if czytest2 == 1:
        if cz == '1':
            ktoraminuta = 'minute'
        else:
            ktoraminuta = 'minutes'
        teksttestu1 = document.add_paragraph()
        teksttestu1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        if medB == 'hydrostatycznie':
            medben = 'water tested'
        elif medB == 'powietrzem pod wodą':
            medben = 'tested underwater with air'
        elif medB == 'azotem pod wodą':
            medben = 'tested underwater with nitrogen'
        if ile == 1:
            hashave = 'has'
        elif ile > 1:
            hashave = 'have'
        teksttestu1.add_run(checkIle2+' '+hashave+' been ' + medben + ' under pressure of ' + cisT + ' bar for ' + cz + ' '+ktoraminuta+'.').font.underline = TRUE

        teksttestu2 = document.add_paragraph()
        teksttestu2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        teksttestu2.add_run('No leakage was observed.').font.underline = TRUE
        document.add_paragraph()
    else:
        pass

    if czyodtlu.get() == 1:
        tekstodtluszczania = document.add_paragraph()
        tekstodtluszczania.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        tekstodtluszczania.add_run(checkIle2+' underwent preliminary degreasing procedure.').font.underline = TRUE
        tekstodtluszczania2 = document.add_paragraph()
        tekstodtluszczania2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        tekstodtluszczania2.add_run('Final degreasing and approval to work with oxygen must be done by customer.').font.underline = TRUE
        document.add_paragraph()
    else:
        pass

    if czyzuzycie.get() == 1:
        if ile == 1:
            tekstzuzycia = document.add_paragraph()
            tekstzuzycia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
            tekstzuzycia.add_run('Due to variety of media and their aggressiveness, service life of the hose assembly is limited.').font.underline = TRUE
            document.add_paragraph()
        else:
            tekstzuzycia = document.add_paragraph()
            tekstzuzycia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
            tekstzuzycia.add_run('Due to variety of media and their aggressiveness, service life of the hose assemblies is limited..').font.underline = TRUE
            document.add_paragraph()
    else:
        pass

    if czyciaglosc.get() == 1:
        tekstciaglosc = document.add_paragraph()
        tekstciaglosc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # @UndefinedVariable
        tekstciaglosc.add_run(
            'Electric conductivity between the fittings has been confirmed.').font.underline = TRUE
        document.add_paragraph()
    else:
        pass

    document.add_paragraph(checkIle2+' can work with parameters stated by the customer.')
    document.add_paragraph()

    uzytkowanie1 = document.add_paragraph(checkIle2+' can work according to technical state. The operator should take into consideration mechanical wear, abrasions, operator risk and all the factor which can affect the assembly usage and lifetime.')
    uzytkowanie1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY # @UndefinedVariable

    document.add_paragraph()

    if ktoryoddzial.get() == '50':
        przygzatw = document.add_paragraph('Prepared by:\t' + przygotowal + '\tApproved by:\t' + zatwierdzil)
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(14))

        datydaty = document.add_paragraph('Date:\t' + kiedy_p + '\tDate:\t' + kiedy_z)
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(14))

        podpispodpis = document.add_paragraph('Signature:\tSignature:')
        podpispodpis.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
        podpispodpis.paragraph_format.tab_stops.add_tab_stop(Cm(11.5))
    else:
        przygzatw = document.add_paragraph('Prepared by:\t' + przygotowal)
        przygzatw.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))

        datydaty = document.add_paragraph('Date:\t' + kiedy_p)
        datydaty.paragraph_format.tab_stops.add_tab_stop(Cm(2.5))

        podpispodpis = document.add_paragraph('Signature:')

    def filesave():
        if ktoryoddzial.get() == '50':
            sss = nQC
        else:
            sss = 'QC-'+ktoryoddzial.get()+'-'+nQC+'-'+strftime('%y')
        filename = filedialog.asksaveasfilename(defaultextension='.docx', initialfile=sss)
        if not filename: return
        document.save(filename)
        pytanieootwarcie = messagebox.askyesno("Co teraz?", "Gratuluję sukcesu!\nCzy chcesz teraz otworzyć zapisany dokument?", default="yes")
        if pytanieootwarcie:
            startfile(filename)
        else:
            pass
    filesave()


class niepustentry(Tk):
    def __init__(self, rodzic, clmn, rw):
        Entry.__init__(self, rodzic)
        vcmd = (rodzic.register(self.validate))
        self.nazwaentry = Entry(rodzic, validate='focus', validatecommand=vcmd)
        self.nazwaentry.grid(padx=2, pady=2, row=rw, sticky=E, column=clmn)

    def validate(self):
        if self.nazwaentry.get() == "":
            self.nazwaentry.config(bg="misty rose")
            return True
        else:
            self.nazwaentry.config(bg="pale green")
            return True

    def get(self):
        return self.nazwaentry.get()

    def delete(self,x,y):
        return self.nazwaentry.delete(x, y)

    def insert(self,x,y):
        return self.nazwaentry.insert(x, y)

    def disable(self):
        return self.nazwaentry.config(state='disabled')

    def activate(self):
        return self.nazwaentry.config(state='normal')

    def focus(self):
        return self.nazwaentry.focus()

    def bind(self, k, f):
        return self.nazwaentry.bind(k, f)

class vacuum(niepustentry):
    def __init__(self, rodzic, clmn, rw):
        super().__init__(rodzic, clmn, rw)

    def validate(self):
        if isfloat(self.nazwaentry.get()) == True and 0 <= float(self.nazwaentry.get().replace(',','.')) <= 1:
            self.nazwaentry.config(bg="pale green")
            return True
        else:
            self.nazwaentry.config(bg="misty rose")
            return True

class liczbaentry(niepustentry):
    def __init__(self, rodzic, clmn, rw):
        super().__init__(rodzic, clmn, rw)

    def validate(self):
        if self.nazwaentry.get() == "" or isfloat(self.nazwaentry.get()) == False:
            self.nazwaentry.config(bg="misty rose")
            return True
        else:
            self.nazwaentry.config(bg="pale green")
            return True

class klikonentry(niepustentry):
    def __init__(self, rodzic, clmn, rw):
        super().__init__(rodzic, clmn, rw)
        self.nazwaentry.config(width=33)

    def validate(self):
        if self.nazwaentry.get() == "":
            self.nazwaentry.config(bg="misty rose")
            return True
        else:
            self.nazwaentry.config(bg="pale green")
            return True


class numerQC1(niepustentry):
    def __init__(self,rodzic, clmn, rw):
        super().__init__(rodzic, clmn, rw)

    def validate(self):
        if self.nazwaentry.get() == "":
            self.nazwaentry.config(bg="white")
            return True
        if self.nazwaentry.get().isdigit() and len(self.nazwaentry.get()) == 4:
            self.nazwaentry.config(bg="pale green")
            return True
        else:
            self.nazwaentry.config(bg="misty rose")
            return True

def isfloat(s):
    try:
        float(s.replace(',', '.'))
        return True
    except ValueError:
        return False

ramkaKlient = LabelFrame(root, text=" Informacje o kliencie ")
ramkaKlient.columnconfigure(1, minsize=160)
ramkaKlient.grid(padx=1, pady=2, ipadx=5, ipady=5, sticky=N+S+E+W, column=0, row=0)

labelKlient1 = Label(ramkaKlient, text="Nazwa klienta:")
labelKlient1.grid(padx=2, pady=2, sticky=W, row=0)

entryKlient1 = klikonentry(ramkaKlient, 1, 0)
entryKlient1.focus()
entryKlient1.bind('<Return>', znajdzklientaenter)

labelKlient2 = Label(ramkaKlient, text="Adres 1:")
labelKlient2.grid(padx=2, pady=2, sticky=W, row=1)

entryKlient2 = klikonentry(ramkaKlient, 1, 1)

labelKlient3 = Label(ramkaKlient, text="Adres 2:")
labelKlient3.grid(padx=2, pady=2, sticky=W, row=2)

entryKlient3 = klikonentry(ramkaKlient, 1, 2)

Button(ramkaKlient, text='Znajdź klienta', command=znajdzklienta).grid(padx=2, pady=2, row=3, column=0, columnspan=2, sticky=W)
przyciskdodajklienta = Button(ramkaKlient, text="Dodaj klienta", command=dodajklienta)
przyciskdodajklienta.grid(padx=2, pady=2, row=3, column=0, columnspan=2, sticky=E)

gornaRamka = LabelFrame(root, text=" Ważne numery ")
gornaRamka.columnconfigure(1, minsize=160)
gornaRamka.grid(padx=1, pady=2, ipadx=5, ipady=3, sticky=N + S + E + W, column=0, row=1)

tekstnumerQC = Label(gornaRamka, text="Numer świadectwa QC")
tekstnumerQC.grid(padx=2, pady=2, sticky=W, row=0)

numerQC = numerQC1(gornaRamka, 1, 0)

tekstnumerKRP = Label(gornaRamka, text="Numer Zlec. Prod.")
tekstnumerKRP.grid(padx=2, pady=2, sticky=W, row=1)

numerZP = niepustentry(gornaRamka, 1, 1)
numerZP.insert(END, '*17012345-1-1')

tekstnumerTU = Label(gornaRamka, text="Numer TU")
tekstnumerTU.grid(padx=2, pady=2, sticky=W, row=2)

tu = niepustentry(gornaRamka, 1, 2)
tu.insert(END, 'TU-PL50-17-01234')

labelZamowienie = Label(gornaRamka, text='Zamówienie')
labelZamowienie.grid(padx=2, pady=2, sticky=W, row=3)

numerZam = Entry(gornaRamka, width=20)
numerZam.grid(padx=2, pady=2, column=1, row=3, sticky=E, columnspan=3)

koncowkaframe = LabelFrame(root, text=' Końcówki przewodu ')
koncowkaframe.grid(padx=1, pady=2, row=2, column=0, sticky=W+E+S+N)

koncowka1label = Label(koncowkaframe, text='Końcówka 1:')
koncowka1label.grid(padx=2, pady=2, sticky=W, column=0, row=1)

koncowka1entry = klikonentry(koncowkaframe, 1, 1)

koncowka2label = Label(koncowkaframe, text='Końcówka 2:')
koncowka2label.grid(padx=2, pady=2, sticky=W, column=0, row=2)

koncowka2entry = klikonentry(koncowkaframe, 1, 2)

def koncowkacheck():
    if ilekoncowek.get() == 2:
        koncowka2entry.disable()
    elif ilekoncowek.get() == 1:
        koncowka2entry.activate()
    else:
        pass

ilekoncowek = IntVar()
dwiekoncowki = Radiobutton(koncowkaframe, text="Dwie różne", variable=ilekoncowek, value=1, command=koncowkacheck)
dwiekoncowki.grid(sticky=W, row=0, column=0)
dwiekoncowki.invoke()

jednakoncowka = Radiobutton(koncowkaframe, text="Takie same", variable=ilekoncowek, value=2, command=koncowkacheck)
jednakoncowka.grid(sticky=E, row=0, column=1)

materialramka = LabelFrame(root, text=" Materiał końcówek i tulei ")
materialramka.grid(padx=1, pady=2, ipadx=5, ipady=5, row=3, column=0, sticky=W+E+N+S)

materialklabel = Label(materialramka, text='Wybierz materiał końcówek:')
materialklabel.grid(padx=2, pady=2, sticky=W, column=0, row=0)

listamaterialow = {
    'stal węglowa': 'ze stali węglowej',
    'stal nierdzewna': 'ze stali nierdzewnej',
    'stal węglowa ocynkowana': 'ze stali węglowej ocynkowanej',
    'aluminium': 'z aluminium',
    'mosiądz': 'z mosiądzu',
    'st. nierdz. AISI 304': 'ze stali nierdzewnej AISI 304',
    'st. nierdz. AISI 316': 'ze stali nierdzewnej AISI 316'
    }

listamaterialowang = {
    'stal węglowa': 'carbon steel',
    'stal nierdzewna': 'stainless steel',
    'stal węglowa ocynkowana': 'zinc-coated carbon steel',
    'aluminium': 'aluminum',
    'mosiądz': 'brass',
    'st. nierdz. AISI 304': 'stainless steel AISI 304',
    'st. nierdz. AISI 316': 'stainless steel AISI 316'
    }


lmlm = ['stal węglowa', 'stal nierdzewna', 'stal węglowa ocynkowana', 'aluminium', 'mosiądz', 'st. nierdz. AISI 304', 'st. nierdz. AISI 316']

Label(materialramka, text='Końcówki:').grid(padx=2, pady=2, row=1, column=0, sticky=W)

koncowki = ttk.Combobox(materialramka, state='readonly', width=20)
koncowki['values'] = lmlm
koncowki.grid(row=1, column=1)
koncowki.current([0])

materialtlabel = Label(materialramka, text='Wybierz materiał tulei/obejm:')
materialtlabel.grid(padx=2, pady=2, sticky=W, column=0, row=2)

tulejwybor = ttk.Combobox(materialramka, state='readonly', width=20)
tulejwybor['values'] = ['tuleje', 'obejmy']
tulejwybor.grid(padx=5, pady=2, row=3, sticky=W, column=0)
tulejwybor.current([0])

tuleje = ttk.Combobox(materialramka, state='readonly', width=20)
tuleje['values'] = lmlm
tuleje.grid(row=3, column=1)
tuleje.current([0])


dolnaRamka = LabelFrame(root, text=" Informacje o przewodzie ")
dolnaRamka.columnconfigure(0, minsize=170)
dolnaRamka.grid(ipadx=2, padx=1, pady=2, row=0, column=1, rowspan=3,sticky=N+S+E+W)

labelIlosc = Label(dolnaRamka, text="Ilość")
labelIlosc.grid(padx=2, pady=2, sticky=W, row=1)

ilosc = liczbaentry(dolnaRamka, 0, 1)

waz = ttk.Combobox(dolnaRamka, state='readonly', width=22)
waz['values'] = ('Długość przewodu (mm)', 'Długość węża (mm)')
waz.current([0])
waz.grid(padx=2, pady=2, row=2, column=0, sticky=W)

dlugosc = liczbaentry(dolnaRamka, 0, 2)

ramkawyboru = LabelFrame(dolnaRamka, text=' Wybór węża ')
ramkawyboru.grid(padx=5, pady=2, ipadx=5, row=3, column=0, sticky=N)

tekstindeksuweza = StringVar()
indeksweza = Label(ramkawyboru, textvariable=tekstindeksuweza, font=('TkDefaultFont', 10, 'bold'))
indeksweza.grid(padx=2, pady=2, sticky=E, column=0, row=4)

labelwprowadzwaz = Label(ramkawyboru, text='Wprowadź indeks:')
labelwprowadzwaz.grid(padx=2, pady=2, sticky=W, column=0, row=1)

wpisanywaz = StringVar()
szukajweza = Entry(ramkawyboru, width=23, textvariable=wpisanywaz)
szukajweza.grid(padx=2, pady=2, sticky=E, column=0, row=1)
wpisanywaz.trace('w', szukaj)

rezultaty = Listbox(ramkawyboru, height=14, activestyle='none', width=45)
rezultaty.grid(row=2, column=0, pady=2, padx=2)
rezultaty.bind("<Double-Button-1>", kliknijwybor)

yscrollbar = Scrollbar(ramkawyboru, orient=VERTICAL)
yscrollbar.config(command=rezultaty.yview)
yscrollbar.grid(row=2, column=0, sticky=N+S+E, pady=4, padx=4)

rezultaty.config(yscrollcommand=yscrollbar.set)

guziczek = Button(ramkawyboru, text='Wybierz wąż', command=wprowadzwybor)
guziczek.grid(ipadx=5,ipady=5, padx=10, pady=8, sticky=W, column=0, row=4)

warunkiKlienta = LabelFrame(root, text=" Warunki robocze klienta ")
warunkiKlienta.grid(padx=1, pady=2, row=0, column=2, sticky=E+W+N+S)

medium = niepustentry(warunkiKlienta, 1, 0)
podcisn = vacuum(warunkiKlienta, 1, 1)
podcisn.disable()
cisnienieRob = niepustentry(warunkiKlienta, 1, 2)
tempWew = niepustentry(warunkiKlienta, 1, 3)
tempZew = niepustentry(warunkiKlienta, 1, 4)

labelMedium = Label(warunkiKlienta, text="Medium")
labelMedium.grid(padx=2, pady=2, sticky=W, row=0, column=0)

labelPodcisnienie = Label(warunkiKlienta, text="Podciśnienie (bar)")
labelPodcisnienie.grid(padx=2, pady=2, sticky=W, row=1, column=0)

def vacuumactivation():
    if czypodcisn.get() == 0:
        podcisn.disable()
    elif czypodcisn.get() == 1:
        podcisn.activate()

czypodcisn = IntVar()
czekpodcisn = Checkbutton(warunkiKlienta, variable=czypodcisn, command=vacuumactivation)
czekpodcisn.grid(padx=2, pady=2, sticky=E, column=0, row=1)
czekpodcisn.deselect()

labelCisRob = Label(warunkiKlienta, text="Ciśnienie robocze (bar)")
labelCisRob.grid(padx=2, pady=2, sticky=W, row=2)

labelTempWew = Label(warunkiKlienta, text="Temp. wewnętrzna (°C)")
labelTempWew.grid(padx=2, pady=2, sticky=W, row=3)

labelTempZew = Label(warunkiKlienta, text="Temp. zewnętrzna (°C)")
labelTempZew.grid(padx=2, pady=2, sticky=W, row=4)

warunkiTestu = LabelFrame(root, text=" Warunki testu ")
warunkiTestu.grid(padx=1, pady=2, sticky=W+E+N+S, row=1, column=2)

czymBadane = Label(warunkiTestu, text="Rodzaj testu")
czymBadane.grid(padx=2, pady=2, sticky=W, row=1, column=0)

mediumBadania = ttk.Combobox(warunkiTestu, state='readonly')
mediumBadania['value'] = ('hydrostatycznie', 'powietrzem pod wodą', 'azotem pod wodą')
mediumBadania.current(0)
mediumBadania.grid(padx=2, pady=2, row=1, column=1)

labelCisTest = Label(warunkiTestu, text="Ciśnienie testu (bar)")
labelCisTest.grid(padx=2, pady=2, sticky=W, row=3, column=0)

cisnienieTest = liczbaentry(warunkiTestu, 1, 3)

labelCzasTestu = Label(warunkiTestu, text="Czas testu (min)")
labelCzasTestu.grid(padx=2, pady=2, sticky=W, row=4, column=0)

czasTestu = liczbaentry(warunkiTestu, 1, 4)

def naccheck():
    if czytest.get() == 2:
        mediumBadania.config(state='disabled')
        cisnienieTest.disable()
        czasTestu.disable()
    elif czytest.get() == 1:
        mediumBadania.config(state='readonly')
        cisnienieTest.activate()
        czasTestu.activate()
    else:
        pass

czytest = IntVar()
badanieTak = Radiobutton(warunkiTestu, text="TAK", variable=czytest, value=1, command=naccheck)
badanieTak.grid(sticky=W, row=0, column=0)

badanieNie = Radiobutton(warunkiTestu, text="NIE", variable=czytest, value=2, command=naccheck)
badanieNie.grid(sticky=W, row=0, column=1)


ktokiedyRamka = LabelFrame(root, text=" Przygotował i zatwierdził ")
ktokiedyRamka.columnconfigure(1, minsize=200)
ktokiedyRamka.grid(padx=1, pady=2, row=2, column=2, sticky=N+S+E+W)

labelprzygotowal = Label(ktokiedyRamka, text='Przygotował:')
labelprzygotowal.grid(padx=2, pady=2, sticky=W, row=0, column=0)

listaprzygotowal = {
    '50': ['Adam Rogacewicz', 'Damian Średzki', 'Arkadiusz Klamerek', 'Mariusz Węcławiak', 'Michał Rosada', 'Robert Tomaszewski', 'Krzysztof Wenda', 'Hubert Stanisławski', 'Daniel Józefiak'],
    '02': ['Krzysztof Szrejter', 'Małgorzata Smug', 'Marek Falas', 'Robert Woźniak'],
    '04': ['Jarosław Walczak', 'Tomasz Reczek'],
    '05': ['Sylwia Pik', 'Wojciech Lejmel', 'Arkadiusz Gajdamowicz', 'Michał Zahorski', 'Piotr Sikora'],
    '07': ['Dariusz Kentop', 'Przemysław Kapela'],
    '15': ['Piotr Muraczewski', 'Paweł Skoczypiec', 'Marcin Trybczyński', 'Joanna Stangreciak'],
    '18': ['Tomasz Grzybek', 'Jakub Neczaj', 'Paweł Pudło'],
    '20': ['Tomasz Napiórkowski', 'Eugeniusz Podolak'],
    '61': ['Dariusz Kentop', 'Przemysław Kapela']
    }

wczytajoddzial()

if wlaczniktest.get() == 1:
    badanieTak.invoke()
else:
    badanieNie.invoke()

entryprzygotowal = ttk.Combobox(ktokiedyRamka, state='readonly')
entryprzygotowal['values'] = listaprzygotowal.get(ktoryoddzial.get())
entryprzygotowal.grid(padx=2, pady=2, sticky=E, column=1, row=0)
entryprzygotowal.current(0)

labelkiedyprzygotowal = Label(ktokiedyRamka, text='Data:')
labelkiedyprzygotowal.grid(padx=2, pady=2, sticky=W, column=0, row=1)

now = datetime.now()
biezacadata = strftime('%d.%m.%Y')

entrykiedyprzygotowal = Entry(ktokiedyRamka)
entrykiedyprzygotowal.insert(END, biezacadata)
entrykiedyprzygotowal.grid(padx=2, pady=2, sticky=E, column=1, row=1)

labelzatwierdzil = Label(ktokiedyRamka, text='Zatwierdził:')
labelzatwierdzil.grid(padx=2, pady=2, sticky=W, row=2, column=0)

entryzatwierdził = ttk.Combobox(ktokiedyRamka, state='readonly')
entryzatwierdził['values'] = ['Mariusz Węcławiak', 'Hubert Stanisławski', 'Krzysztof Wenda', 'Daniel Józefiak']
entryzatwierdził.grid(padx=2, pady=2, sticky=E, column=1, row=2)
entryzatwierdził.current(0)

labelkiedyzatwierdzil = Label(ktokiedyRamka, text='Data:')
labelkiedyzatwierdzil.grid(padx=2, pady=2, sticky=W, column=0, row=3)

entrykiedyzatwierdzil = Entry(ktokiedyRamka)
entrykiedyzatwierdzil.insert(END, biezacadata)
entrykiedyzatwierdzil.grid(padx=2, pady=2, sticky=E, column=1, row=3)


ramkauwagi = LabelFrame(root, text=' Uwagi / adnotacje (pod informacją o końcówkach): ')
ramkauwagi.grid(padx=1, pady=2, row=3, column=1, sticky=W+E+S+N)

poleuwagi = Text(ramkauwagi, height=6, width=34)
poleuwagi.grid(padx=9, pady=5)

scrolluwagi = Scrollbar(ramkauwagi, orient=VERTICAL)
scrolluwagi.config(command=poleuwagi.yview)
scrolluwagi.grid(row=0, column=0, sticky=N+S+E, pady=5)

poleuwagi.config(yscrollcommand=scrolluwagi.set)


ramkaoznaczenia = LabelFrame(root, text=' Niestandardowe oznaczenie ')
ramkaoznaczenia.grid(padx=1, pady=2, row=4, column=1, sticky=E+W+N+S)

typoznaczenia = ttk.Combobox(ramkaoznaczenia, state='readonly', width=45)
typoznaczenia['values'] = ['NIE - oznaczenie standardowe', 'Pharmaline N/G', 'Spir Star', 'CORROFLON', 'CE I (stalowy)', 'CE II (stalowy)', 'COLGATE', 'CE I (niestalowy)', 'PESA (OEM)']
typoznaczenia.grid(padx=4, pady=2, column=0, row=0)
typoznaczenia.bind('<<ComboboxSelected>>', onselectoznakowanie)
typoznaczenia.current(0)

poleoznaczenia = Text(ramkaoznaczenia, width=35, height=4)
poleoznaczenia.grid(padx=4, pady=5, column=0, row=1)


dodatkoweinformacje = LabelFrame(root, text=' Dodatkowe informacje ')
dodatkoweinformacje.columnconfigure(0, minsize=270)
dodatkoweinformacje.grid(padx=1, pady=2, row=3, column=2, sticky=E+W+N+S)

czyodtlu = IntVar()
odtluszczane = Checkbutton(dodatkoweinformacje, text='Notka o odtłuszczaniu', variable=czyodtlu)
odtluszczane.grid(padx=2, pady=5, sticky=W, column=0, row=0)

czyzuzycie = IntVar()
ograniczonezycie = Checkbutton(dodatkoweinformacje, text='Notatka o ograniczonej żywotności', variable=czyzuzycie)
ograniczonezycie.grid(padx=2, pady=5, sticky=W, column=0, row=1)

czyciaglosc = IntVar()
ciagloscelektryczna = Checkbutton(dodatkoweinformacje, text='Notatka o ciągłości elektrycznej przewodu', variable=czyciaglosc)
ciagloscelektryczna.grid(padx=2, pady=5, sticky=W, column=0, row=2)

przyciskisterowania = LabelFrame(root, text=' Tworzenie dokumentu ')
przyciskisterowania.grid(padx=1, pady=2, row=4, column=2, sticky=E+W+N+S)

wersjajezykowa = StringVar()

def wlaczPL():
    if wersjajezykowa.get() == 'PL':
        pass
    else:
        wersjajezykowa.set('PL')
        jezykpl.config(image=img_pl_on)
        jezykpl.config(relief=RAISED)
        jezykeng.config(image=img_eng_off)
        jezykeng.config(relief=FLAT)
        jezykeng.config(overrelief=RAISED)
        przyciskGeneruj.config(text="ZAPISZ\nŚWIADECTWO\nPL")
        szukajweza.delete(0, END)
        rezultaty.delete(0, END)
        tekstindeksuweza.set('')
        ladujlistewezy('PL')

def wlaczENG():
    if wersjajezykowa.get() == 'PL':
        wersjajezykowa.set('ENG')
        jezykpl.config(image=img_pl_off)
        jezykpl.config(relief=FLAT)
        jezykpl.config(overrelief=RAISED)
        jezykeng.config(image=img_eng_on)
        jezykeng.config(relief=RAISED)
        przyciskGeneruj.config(text="ZAPISZ\nŚWIADECTWO\nENG")
        szukajweza.delete(0, END)
        rezultaty.delete(0, END)
        tekstindeksuweza.set('')
        ladujlistewezy('ENG')
    else:
        pass

pl_on = Image.open('bin/pl_on.png')
img_pl_on = PhotoImage(pl_on)

pl_off = Image.open('bin/pl_off.png')
img_pl_off = PhotoImage(pl_off)

eng_off = Image.open('bin/ang_off.png')
img_eng_off = PhotoImage(eng_off)

eng_on = Image.open('bin/ang_on.png')
img_eng_on = PhotoImage(eng_on)

jezykpl = Button(przyciskisterowania, command=wlaczPL)
jezykpl.grid(column=0, row=0, padx=(10,0))

jezykeng = Button(przyciskisterowania, command=wlaczENG)
jezykeng.grid(column=1, row=0)

przyciskGeneruj = Button(przyciskisterowania, text="ZAPISZ\nŚWIADECTWO", height=4, width=15, command=sprawdzenieprzedzapisem)
przyciskGeneruj.grid(padx=30, pady=20, column=2, row=0)

wlaczPL()

def wklejzezlecenia():
    s = pyperclip.paste()
    if "$LU=ShopOrd" in s:
        for i in s.splitlines():
            if "$3:PART_NO" in i:
                _, b = i.split("=")
                tu.delete(0, END)
                tu.insert(END, b)
                tu.validate()
            if "$0:ORDER_NO" in i:
                _, b = i.split("=")
                numerZP.delete(0, END)
                numerZP.insert(END, b)
            if "$1:RELEASE_NO" in i:
                _, b = i.split("=")
                numerZP.insert(END, "-")
                numerZP.insert(END, b)
            if "$2:SEQUENCE_NO" in i:
                _, b = i.split("=")
                numerZP.insert(END, "-")
                numerZP.insert(END, b)
                numerZP.validate()
            if "$9:REVISED_QTY_DUE" in i:
                _, b = i.split("=")
                ilosc.delete(0, END)
                ilosc.insert(END, b)
                ilosc.validate()
            if "$127:CUST_ORD_CUSTOMER_API.GET_NAME(C_CUSTOMER_NO)" in i:
                _, b = i.split("=")
                entryKlient1.delete(0, END)
                entryKlient2.delete(0, END)
                entryKlient3.delete(0, END)
                entryKlient1.insert(END, b)
                znajdzklienta()
    else:
        messagebox.showinfo("Błąd!", "Upewnij się czy skopiowano poprawny obiekt IFS!")

def wklejzkartyinf():
    s = pyperclip.paste()
    if "$LU=CInfCardDet" in s:
        for i in s.splitlines():
            if "$3:C_WORK_PREASSURE" in i:
                _, b = i.split("=")
                cisnienieRob.delete(0, END)
                cisnienieRob.insert(END, b)
                cisnienieRob.validate()
            if "$4:C_WORK_UNDERPREASSURE" in i:
                _, b = i.split("=")
                if b == "" or b == "0":
                    podcisn.activate()
                    podcisn.delete(0, END)
                    czypodcisn.set(0)
                    podcisn.disable()
                else:
                    podcisn.activate()
                    czypodcisn.set(1)
                    podcisn.delete(0, END)
                    podcisn.insert(END, b)
                    podcisn.validate()
            if "$5:C_INT_WORK_TEMP" in i:
                _, b = i.split("=")
                if b.isdigit():
                    if int(b) > 0:
                        b = '%+d' % int(b)
                tempWew.delete(0, END)
                tempWew.insert(END, b)
                tempWew.validate()
            if "$6:C_EXT_WORK_TEMP" in i:
                _, b = i.split("=")
                if b.isdigit():
                    if int(b) > 0:
                        b = '%+d' % int(b)
                tempZew.delete(0, END)
                tempZew.insert(END, b)
                tempZew.validate()
            if "$7:C_MEDIUM" in i:
                _, b = i.split("=")
                if b == "999":
                    for i in s.splitlines():
                        if "$14:C_MEDIUM_ADD_INFO" in i:
                            _, b = i.split("=")
                            if "(gazowy)" in b:
                                b = b.replace('(gazowy)', '')
                            medium.delete(0, END)
                            medium.insert(END, b)
                            medium.validate()
                else:
                    for i in s.splitlines():
                        if "$13:C_MEDIUM_API.GET_DESCRIPTION (C_MEDIUM)" in i:
                            _, b = i.split("=")
                            if "(gazowy)" in b:
                                b = b.replace(' (gazowy)', '')
                            medium.delete(0, END)
                            medium.insert(END, b)
                            medium.validate()
            if "$10:C_TUBE_LENGHT" in i:
                _, Lw = i.split("=")
                if Lw == "" or Lw == "0":
                    for i in s.splitlines():
                        if "$9:C_TOTAL_LENGHT" in i:
                            _, b = i.split("=")
                            dlugosc.delete(0, END)
                            dlugosc.insert(END, b)
                            waz.current(0)
                            dlugosc.validate()
                else:
                    dlugosc.delete(0, END)
                    dlugosc.insert(END, Lw)
                    waz.current(1)
                    dlugosc.validate()
    else:
        messagebox.showinfo("Błąd!", "Upewnij się czy skopiowano poprawny obiekt IFS!")


slownikoddzialow = {
    "01": "Eksport",
    "02": "Kalisz",
    "03": "Łódź",
    "04": "Katowice",
    "05": "Gdańsk",
    "06": "Gdynia",
    "07": "Poznań",
    "08": "Warszawa",
    "09": "Gliwice",
    "10": "Kraków",
    "11": "Białystok",
    "12": "Rzeszów",
    "13": "Wrocław",
    "14": "Szczecin",
    "15": "Toruń",
    "16": "Lublin",
    "17": "Kielce",
    "18": "Częstochowa",
    "19": "Zielona Góra",
    "20": "Olsztyn",
    "21": "Bydgoszcz",
    "22": "Płock",
    "23": "Bielsko-Biała",
    "24": "Przeźmierowo",
    "50": "Centrala",
    "61": "Produkcja OEM"
    }


root.title("Świadectwomat v"+str(versionnr)+" - "+slownikoddzialow.get(ktoryoddzial.get()))

def liniadlaexcela():
    ntu = tu.get()
    if "CZ" in ntu:
        oddzial = "Czechy"
    elif "SK" in ntu:
        oddzial = "Słowacja"
    elif "LT" in ntu:
        oddzial = "Litwa"
    elif "RU" in ntu:
        oddzial = "Rosja"
    elif "PL" in ntu:
        nroddzialu = "".join([ntu[5], ntu[6]])
        if nroddzialu.isdigit():
            oddzial = slownikoddzialow.get(nroddzialu)
        else:
            oddzial = ""
    else:
        oddzial = ""
    if ktoryoddzial.get() == '50':
        pyperclip.copy(strftime('%d-%m-%Y')+"\t"+entryprzygotowal.get()+"\tJ\t"+entryKlient1.get()+"\t"+tekstindeksuweza.get()+"\t"+medium.get()+"\t"+oddzial+"\t"+tu.get()+"\t"+numerZP.get())
    else:
        pyperclip.copy(strftime('%d-%m-%Y')+"\t"+entryprzygotowal.get()+"\tJ\t"+entryKlient1.get()+"\t"+tekstindeksuweza.get()+"\t"+medium.get()+"\t"+tu.get()+"\t"+numerZP.get())

def zglosblad():
    pytanko = messagebox.askyesno("Czy otworzyć okno email?", "Zamierasz wysłać email ze zgłoszeniem błędu do autora programu.\nCzy chcesz kontynuować?", default="no")
    if pytanko:
        url = 'mailto://Mariusz Węcławiak <m.weclawiak@tubes-international.com>?subject=Zgłoszenie błędu w Świadectwomacie'
        webbrowser.open(url,new=1)

guziczki = LabelFrame(root, text=" Przyciski pomocnicze ")
guziczki.grid(padx=1, pady=2, row=4, column=0, sticky=E+W+S+N)

dodajnowywaz = Button(guziczki, text='Dodaj\nnowy wąż', command=otworznowywaz, height=2, width=10)
dodajnowywaz.grid(row=1, column=0, padx=10, pady=10)

guzikwklejzezlecenia = Button(guziczki, text="Wklej\nZlec. Prod.", command=wklejzezlecenia, height=2, width=10)
guzikwklejzezlecenia.grid(sticky=N+S, column=0, row=0, padx=10, pady=2)

guzikwklejzkartyinf = Button(guziczki, text="Wklej\nKartę Inf.", command=wklejzkartyinf, height=2, width=10)
guzikwklejzkartyinf.grid(sticky=N+S, column=1, row=0, padx=10, pady=2)

guzikdoschowka = Button(guziczki, text="Skopiuj linię\ndla Excela", height=2, width=10, command=liniadlaexcela)
guzikdoschowka.grid(column=2, row=0, padx=10, pady=2, sticky=N+S)

przyciskwyczysc = Button(guziczki, text="Wyczyść\nformularz", height=2, width=10, command=wyczyscpola)
przyciskwyczysc.grid(padx=10, pady=10, column=1, row=1, sticky=N+S)


wyslijmaila = Button(guziczki, text='Zgłoś błąd', height=2, width=10, command=zglosblad)
wyslijmaila.grid(padx=10, pady=10, column=2, row=1, sticky=N+S)

kimjestes()
checkupdate()


Tooltip(entryKlient1, text='Wprowadź fragment lub całą nazwę klienta. Wyszukaj klawiszem Enter lub klikając na przycisk "Znajdź klienta" poniżej')
Tooltip(entryKlient2, text='Pole adresu klienta - ulicy z numerem')
Tooltip(entryKlient3, text='Pole adresu klienta - kod pocztowy, miasto')
Tooltip(przyciskdodajklienta, text='Po kliknięciu dane klienta wpisane powyżej zostaną dodane do bazy klientów. Dzięki temu możliwe będzie wczytanie danych klienta przy następnej okazji')
Tooltip(numerQC, text='Numer świadectwa w formacie czterocyfrowym np. 0145, 4521')
Tooltip(numerZP, text='Numer Zlecenia produkcyjnego np. *18015247-2-1')
Tooltip(tu, text='Pełny numer TU')
Tooltip(numerZam, text='Numer zamówienia KLIENTA, pole nieobowiązkowe')
Tooltip(dwiekoncowki, text='Tekst pojawi się w formie:\rKońcówka 1: .....\rKońcówka 2: .....')
Tooltip(jednakoncowka, text='Tekst pojawi się w formie:\rzakuto obustronnie końcówkami .....')
Tooltip(koncowka1entry, text='Opis końcówki np. z GW 1/2" BSPT\rW przypadku tej samej końcówki obustronnie np. z GZ M22x1,5, wersja lekka')
Tooltip(koncowka2entry, text='Opis końcówki np. z GW Rd 65x1/6" wg DIN 11851')
Tooltip(ilosc, text='Ilość przewodów')
Tooltip(dlugosc, text='Długość przewodu lub węża (do wyboru z listy) wyrażona w milimetrach')
Tooltip(szukajweza, text='Fragment szukanego numeru pozycji węża')
Tooltip(rezultaty, text='Pierwsza pozycja listy jest automatycznie wybierana. Wybierz inną dwukrotnym kliknięciem myszki lub przy pomocy przycisku "Wybierz wąż" poniżej')
Tooltip(indeksweza, text='Wybrany wąż')
Tooltip(poleuwagi, text='Pole uwag pod opisem przewodu. Np. "Na długości przewodów zamontowano osłonę PYROJACKET, wykonaną z włókna szklanego pokrytego silikonem"\r"Do każdego przewodu dołączono uszczelkę 15x7x2 mm wykonaną z PTFE"\rlub inne uwagi dotyczące wykonania przewodów')
Tooltip(medium, text='Medium deklarowane przez klienta')
Tooltip(podcisn, text='Deklarowane podciśnienie w formacie od 0 bar podciśnienia (ciśnienie atmosferyczne) do 1 bar podciśnienia (próżnia absolutna)')
Tooltip(cisnienieRob, text='Deklarowane ciśnienie pracy')
Tooltip(tempWew, text='Deklarowana temperatura wewnętrzna w pełnych stopniach. Symbol "°C" zostanie dodany automatycznie. Dla dodatniej wartości znak "+" zostanie dodany automatycznie')
Tooltip(tempZew, text='Deklarowana temperatura zewnętrzna w pełnych stopniach. Symbol "°C" zostanie dodany automatycznie. Dla dodatniej wartości znak "+" zostanie dodany automatycznie')
Tooltip(badanieTak, text='Czy zostało przeprowadzona próba ciśnieniowa?')
Tooltip(badanieNie, text='Czy została przeprowadzona próba ciśnieniowa?')
Tooltip(mediumBadania, text='Wybór metody testu ciśnieniowego')
Tooltip(cisnienieTest, text='Wartość ciśnienia testu w barach')
Tooltip(czasTestu, text='Długość trwania testu w minutach')
Tooltip(entryprzygotowal, text='Osoba, której nazwisko znajdzie się na świadectwie w polu "Przygotował"')
Tooltip(ograniczonezycie, text='Dodatkowa adnotacja "Ze względu na różnorodność mediów oraz ich agresywność, żywotność przewodów może być ograniczona."')
Tooltip(ciagloscelektryczna, text='Dodatkowa adnotacja "Potwierdzono badaniem zachowanie ciągłości elektrycznej pomiędzy końcówkami przewodów."')
Tooltip(jezykpl, text='Przełącz pomiędzy wersją polską i angielską świadectwa.\rUWAGA! Angielska lista węży jest skromniejsza!')
Tooltip(jezykeng, text='Przełącz pomiędzy wersją polską i angielską świadectwa.\rUWAGA! Angielska lista węży jest skromniejsza!')
Tooltip(guzikwklejzezlecenia, text='W nagłówku zlecenia produkcyjnego, prawym przyciskiem myszy rozwiń menu kontekstowe, z którego wybierz "Edytuj" -> "Kopiuj obiekt". Następnie naciśnij ten przycisk. Do pól programu zostaną wklejone podstawowe dane ze zlecenia.')
Tooltip(guzikwklejzkartyinf, text='W podpiętym do zlecenia oknie Zgłoszenie, przejdź do zakładki Karta informacyjna. W jej polu (NIE w nagłówku okna), prawym przyciskiem myszy rozwiń menu kontekstowe, z którego wybierz "Edytuj" -> "Kopiuj obiekt". Następnie naciśnij ten przycisk. Do pól programu zostaną wklejone dane z karty informacyjnej.')
Tooltip(guzikdoschowka, text='Po wpisaniu wszystkich danych potrzebnych do wystawienia świadectwa, ten przycisk umieszcza w schowku informacje ze zlecenia, które "żywcem" można przekleić do arkuszu Excela ze spisem świadectw.')
Tooltip(wyslijmaila, text='Przycisk otworzy domyślnego klienta poczty, z poziomu którego możliwe będzoe wysłanie maila do autora programu. Oprócz błędów, wszelkie sugestie mile widziane.')


def setzatwierdzil(i):
    if i == '50':
        entryzatwierdził.config(state='readonly')
        entrykiedyzatwierdzil.config(state='normal')
        odtluszczane.config(state='normal')
        dodajnowywaz.config(state='normal')
    else:
        entryzatwierdził.config(state='disabled')
        entrykiedyzatwierdzil.config(state='disabled')
        odtluszczane.config(state='disabled')
        dodajnowywaz.config(state='disabled')

setzatwierdzil(ktoryoddzial.get())

root.mainloop()
